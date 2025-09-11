import os
import copy

from docx import Document

from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table
from docx.table import _Row
from docx.table import _Cell

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class DocxRedactor:
    """
    Класс обработки .docx файлов
    """
    def __init__(self):
        """template_filename: str - имя файла-шаблона
        """
        self.filename : str = None
        self.doc : Document = None
        self._saved = True


    def open(self, filename: str) -> Document:
        """Метод открывает документ для редактирования. К этому
        документу будут применяться вызываемые методы

        Args:
            filename (str): имя документа, который редактируется
        """
        self.doc = Document(filename)
        self.filename = filename
        return self.doc


    def close(self):
        """Закрывает редактируемый документ
        """
        self.doc = None
        self.filename = None


    def save(self):
        """Сохораняет в документ изменения
        """
        self.doc.save(self.filename)
        self._saved = True


    def clone_file(self, src_filename: str, clone_filename: str = None) -> None:
        """
        Метод создаёт копию исходного документа
        """

        src_doc = Document(src_filename)

        output_filename = "cloned.docx" if clone_filename is None else clone_filename

        # Если файл с таким именем уже существует, то удаляем его
        if os.path.exists(output_filename):
            os.remove(output_filename)

        src_doc.save(output_filename)


    def find_paragraph_contains_text(self, search_text: str) -> int:
        """Находит первый параграф, который содержит искомый текст.
        Параграф не обязательно должен состоять только из этого текста

        Args:
            search_text (str): текст, который содержится в параграфе

        Returns:
            Возвращает первый параграф, который содержит искомыйй текст.
            Если такого параграфа нет, то возвращает None
        """
        for index, paragraph in enumerate(self.doc.paragraphs):
            if search_text in paragraph.text:
                return index
        return None  # Not found


    def find_paragraph_consists_of_text(self, search_text: str) -> Paragraph:
        """Находит первый параграф, который состоит только из искомого текста.

        Args:
            search_text (str): текст, из которого состоит параграф

        Returns:
            Возвращает первый параграф, который состоит только из искомого текста.
            Если такого параграфа нет, то возвращает None
        """
        for index, paragraph in enumerate(self.doc.paragraphs):
            if search_text == paragraph.text:
                return index
        return None  # Not found


    def replace_text_in_run(self, run: Run, old_text: str, new_text: str) -> bool:
        """Находит текст в объекте Run и заменяет его на новый.
        Returns:
            Возвращает True, если заменил текст, и False, если нет
        """
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
        return False


    def replace_text_in_paragraph(self, paragraph: Paragraph, old_text: str, new_text: str) -> bool:
        """Находит текст в параграфе и заменяет его на новый
        Returns:
            Возвращает True, если заменил текст хотя бы один раз, и False, если нет
        """
        result = False

        # runs = []
        for run in paragraph.runs:
            # runs.append(run.text)
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                result = True
        # print(f"runs={runs}")

        return result


    def replace_text_in_document(self, old_text: str, new_text: str) -> bool:
        """Находит текст во всём документе и заменяет его на новый
        Returns:
            Возвращает True, если заменил текст хотя бы один раз, и False, если нет
        """
        result = False

        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    result = True

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                result = True

        return result


    def replace_text_in_table_cell(self, cell: _Cell, old_text: str, new_text: str) -> bool:
        """Находит текст ячейке таблицы и заменяет его на новый
        Returns:
            Возвращает True, если заменил текст хотя бы один раз, и False, если нет
        """
        result = False

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    result = True

        return result


    def insert_paragraph(self, index: int = None, text: str = None) -> Paragraph:
        """Создает новый параграф в документе и ставит его на позицию index

        Args:
            index (int, optional): позиция, куда ставить параграф. Defaults to None.
            text (str, optional): текст параграфа. Defaults to None.

        Returns:
            Paragraph: Возвращает параграф, который был создан и вставлен в нужное место
        """
        if text is None:
            new_paragraph = self.doc.add_paragraph()
        else:
            new_paragraph = self.doc.add_paragraph(text)

        _index = len(self.doc.paragraphs) if index is None else index

        # Move the new paragraph to the desired position
        # This step moves it by modifying the internal XML elements
        p = new_paragraph._element
        if _index < len(self.doc.paragraphs):
            target = self.doc.paragraphs[_index]._element
            target.addprevious(p)
        else:
            # If index is at the end, just append it
            self.doc._element.body.append(p)

        return new_paragraph


    def clone_paragraph(self, source_paragraph: Paragraph, target_paragraph: Paragraph):
        """Делает полную копию параграфа со всеми стилями и текстом

        Args:
            source_paragraph (Paragraph): параграф, который нужно скопировать
            target_paragraph (Paragraph): изменяемый параграф
        """
        # index_from = src_index
        # index_to = index_from if dst_index is None else dst_index - 1
        # source_paragraph = doc.paragraphs[index_from]
        # new_paragraph = doc.add_paragraph()

        # Copy paragraph styles
        self.copy_paragraph_styles(source_paragraph, target_paragraph)

        # Copy runs with text and styles
        for run in source_paragraph.runs:
            new_run = target_paragraph.add_run(run.text)
            self.copy_run_styles(run, new_run)


    def copy_paragraph_styles(self, source_paragraph: Paragraph, target_paragraph: Paragraph):
        """Делает полную копию стилей параграфа

        Args:
            source_paragraph (Paragraph): параграф, стили которого нужно скопировать
            target_paragraph (Paragraph): изменяемый параграф
        """
        # Copy paragraph formatting
        source_format = source_paragraph.paragraph_format
        target_format = target_paragraph.paragraph_format

        # Копируем отступы
        if source_paragraph.paragraph_format.left_indent is not None:
            target_paragraph.paragraph_format.left_indent = source_paragraph.paragraph_format.left_indent

        if source_paragraph.paragraph_format.right_indent is not None:
            target_paragraph.paragraph_format.right_indent = source_paragraph.paragraph_format.right_indent

        if source_paragraph.paragraph_format.first_line_indent is not None:
            target_paragraph.paragraph_format.first_line_indent = source_paragraph.paragraph_format.first_line_indent

        # Копируем интервалы
        if source_paragraph.paragraph_format.space_before is not None:
            target_paragraph.paragraph_format.space_before = source_paragraph.paragraph_format.space_before

        if source_paragraph.paragraph_format.space_after is not None:
            target_paragraph.paragraph_format.space_after = source_paragraph.paragraph_format.space_after

        # Копируем межстрочный интервал
        target_paragraph.paragraph_format.line_spacing = source_paragraph.paragraph_format.line_spacing
        target_paragraph.paragraph_format.line_spacing_rule = source_paragraph.paragraph_format.line_spacing_rule

        # Копируем выравнивание
        target_paragraph.alignment = source_paragraph.alignment

        # Копируем стиль (если есть)
        if source_paragraph.style:
            target_paragraph.style = source_paragraph.style

        # target_paragraph.style = source_paragraph.style

        if source_paragraph._element.pPr is not None and source_paragraph._element.pPr.numPr is not None:
            num_id = source_paragraph._element.pPr.numPr.numId.val
            ilvl = source_paragraph._element.pPr.numPr.ilvl.val

            pPr = target_paragraph._element.get_or_add_pPr()
            numPr = pPr.get_or_add_numPr()
            numPr._add_numId().val = num_id
            numPr._add_ilvl().val = ilvl

        # Copy paragraph formatting attributes
        target_format.alignment = source_format.alignment
        target_format.left_indent = source_format.left_indent
        target_format.right_indent = source_format.right_indent
        target_format.space_before = source_format.space_before
        target_format.space_after = source_format.space_after
        target_format.line_spacing = source_format.line_spacing
        target_format.line_spacing_rule = source_format.line_spacing_rule
        target_format.first_line_indent = source_format.first_line_indent
        target_format.widow_control = source_format.widow_control
        target_format.keep_together = source_format.keep_together
        target_format.keep_with_next = source_format.keep_with_next

        # Copy style
        target_paragraph.style = source_paragraph.style


    def copy_run_styles(self, source_run: Run, target_run: Run):
        """Делает полную копию стилей объекта Run

        Args:
            source_paragraph (Paragraph): объект Run, стили которого нужно скопировать
            target_paragraph (Paragraph): изменяемый объект Run
        """
        # Copy run formatting
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        # target_run.strike = source_run.strike
        target_run.font.strike = source_run.font.strike
        target_run.font.size = source_run.font.size
        target_run.font.name = source_run.font.name
        target_run.font.highlight_color = source_run.font.highlight_color
        target_run.font.all_caps = source_run.font.all_caps
        target_run.font.small_caps = source_run.font.small_caps
        target_run.font.subscript = source_run.font.subscript
        target_run.font.superscript = source_run.font.superscript
        target_run.font.color.rgb = source_run.font.color.rgb if source_run.font.color else None


    def paragraph_text_set_bold(self, target_paragraph: Paragraph):
        for run in target_paragraph.runs:
            # run.bold = True
            self.run_text_set_bold(run)


    def run_text_set_bold(self, target_run: Run):
        target_run.bold = True


    def delete_paragraph(self, index: int) -> None:
        """
        Удаляет параграф из документа по заданному индексу.

        Args:
            index (int): Индекс параграфа, который нужно удалить.

        Raises:
            IndexError: Если индекс выходит за пределы допустимого диапазона.
            ValueError: Если попытаться удалить последний оставшийся параграф
                        (документ должен содержать хотя бы один параграф).
        """
        paragraphs = self.doc.paragraphs
        if not 0 <= index < len(paragraphs):
            raise IndexError(f"Index {index} is out of range for {len(paragraphs)} paragraphs.")

        # Нельзя удалять последний параграф — документ должен содержать хотя бы один
        if len(paragraphs) == 1:
            raise ValueError("Cannot delete the last paragraph in the document. A Word document must have at least one paragraph.")

        # Получаем XML-элемент параграфа
        p = paragraphs[index]._element

        # Удаляем его из родительского элемента (тела документа)
        p.getparent().remove(p)


    def insert_row_in_table(self, table: Table, index: int = None) -> _Row:
        """Создаёт и вставляет пустую строку в таблицу.

        Args:
            table (Table): Таблица, в которую надо вставить новую строку
            index (int, optional): Позиция, куда ставим новую строку.
            Если не передать этот параметр, то вставка осуществляется в
            конец таблицы. Defaults to None.

        Returns:
            _Row: Возвращает вставленную строку
        """
        tbl = table._tbl  # Get the underlying CT_Tbl object
        new_tr = tbl.add_tr() # Create a new row XML element

        # Copy the cell structure from an existing row (e.g., the first row)
        # This ensures the new row has the correct number of cells and formatting
        if table.rows:
            first_row_cells = table.rows[0].cells
            for cell in first_row_cells:
                new_tc = new_tr.add_tc()
                # You might need to copy cell properties like width, etc.
                # For simplicity, this example just adds empty cells
        else: # If table is empty, add default cells
            for _ in range(len(table.columns)):
                new_tr.add_tc()

        if index < len(tbl.tr_lst):
            # Insert the new row before the row at the target index
            tbl.tr_lst[index].addprevious(new_tr)
        else:
            # If index is beyond existing rows, append to the end
            tbl.append(new_tr)

        return table.rows[index]


    def clone_table_row(self, source_row: _Row, target_row: _Row):
        """Создаёт полную копию строки с её стилями и текстом

        Args:
            source_row (_Row): строка, которую надо скопировать
            target_row (_Row): изменяемаю строка
        """
        # Copy row properties (like height)
        source_tr_pr = source_row._tr.trPr
        if source_tr_pr is not None:
            # Remove existing properties in target row
            pass
            # target_tr = target_row._tr
            # old_tr_pr = target_tr.trPr
            # if old_tr_pr is not None:
            #     target_tr.remove(old_tr_pr)
            # # Add a copy of the source's trPr
            # target_tr.insert(0, source_tr_pr)

        # Копируем высоту строки
        target_row.height = source_row.height
        target_row.height_rule = source_row.height_rule

        for source_cell, target_cell in zip(source_row.cells, target_row.cells):
            # Copy cell properties (like background color)
            self.copy_cell_properties(source_cell, target_cell)

            # source_tc_pr = source_cell._tc.tcPr
            # if source_tc_pr is not None:
            #     target_tc = target_cell._tc
            #     old_tc_pr = target_tc.tcPr
            #     if old_tc_pr is not None:
            #         target_tc.remove(old_tc_pr)
            #     target_tc.insert(0, source_tc_pr)

            # Copy paragraph and run styles (font, bold, italic, etc.)
            for source_paragraph, target_paragraph in zip(source_cell.paragraphs, target_cell.paragraphs):
                target_paragraph.alignment = source_paragraph.alignment
                self.copy_paragraph_styles(source_paragraph, target_paragraph)

                # Copy runs with text and styles
                for run in source_paragraph.runs:
                    new_run = target_paragraph.add_run(run.text)
                    self.copy_run_styles(run, new_run)


    def copy_row_properties(self, source_row: _Row, target_row: _Row):
        # Копируем высоту строки
        target_row.height = source_row.height
        target_row.height_rule = source_row.height_rule

        for source_cell, target_cell in zip(source_row.cells, target_row.cells):
            self.copy_cell_properties(source_cell, target_cell)


    def copy_cell_properties(self, source_cell: _Cell, target_cell: _Cell):
        """
        Полностью копирует свойства ячейки:
        - Границы (w:tcBorders)
        - Вертикальное выравнивание (w:vAlign)
        - Заливку (w:shd)
        - Подсветку и другие настройки в w:tcPr
        """
        # Получаем исходные и целевые XML-элементы
        src_tc = source_cell._tc
        dst_tc = target_cell._tc
        src_tcPr = src_tc.get_or_add_tcPr()  # w:tcPr исходной ячейки
        dst_tcPr = dst_tc.get_or_add_tcPr()  # w:tcPr целевой ячейки

        # # --- Шаг 1: Очищаем существующие свойства у целевой ячейки ---
        # for child in list(dst_tcPr):
        #     dst_tcPr.remove(child)

        # # Копирование стилей (например, шрифта)
        # if source_cell._element.p_lst:
        #     for p in source_cell._element.p_lst:
        #         target_cell._element.append(p)

        # --- Шаг 2: Копируем нужные элементы из исходного tcPr ---
        tags_to_copy = ['w:tcBorders', 'w:vAlign', 'w:shd', 'w:highlight']

        # self.copy_cell_width(source_cell, target_cell)
        # Копируем ширину ячейки
        if source_cell.width:
            target_cell.width = source_cell.width

        # Копируем вертикальное выравнивание (если задано явно)
        if source_cell.vertical_alignment:
            target_cell.vertical_alignment = source_cell.vertical_alignment

        for tag in tags_to_copy:
            element = src_tcPr.first_child_found_in(tag)
            if element is not None:
                # Глубокая копия XML-узла
                copied_element = copy.deepcopy(element)
                dst_tcPr.append(copied_element)


    def set_border_to_cell(self, cell: _Cell, side: str, value: str, size: int, space: int, color: str = "000000"):
        """
        Безопасно устанавливает границы ячейки без ошибок xmlns.
        value - тип линии
        size - толщина
        space - отступ
        color - цвет
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Удаляем старые границы
        existing_borders = tcPr.first_child_found_in('w:tcBorders')
        if existing_borders is not None:
            tcPr.remove(existing_borders)

        # Создаём новый элемент <w:tcBorders>
        tcBorders = OxmlElement('w:tcBorders')

        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), value)      # тип линии
        border.set(qn('w:sz'), str(size))         # толщина: 4 = 0.5 pt
        border.set(qn('w:space'), str(space))       # отступ
        border.set(qn('w:color'), color)   # цвет
        tcBorders.append(border)

        # Добавляем в ячейку
        tcPr.append(tcBorders)


    def _set_borders_to_cell(self, cell: _Cell):
        """
        Безопасно устанавливает границы ячейки без ошибок xmlns.
        value - тип линии
        size - толщина
        space - отступ
        color - цвет
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Удаляем старые границы
        existing_borders = tcPr.first_child_found_in('w:tcBorders')
        if existing_borders is not None:
            tcPr.remove(existing_borders)

        # Создаём новый элемент <w:tcBorders>
        tcBorders = OxmlElement('w:tcBorders')

        # Для каждой стороны
        for side in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')   # тип линии
            border.set(qn('w:sz'), '4')         # толщина: 4 = 0.5 pt
            border.set(qn('w:space'), '0')      # отступ
            border.set(qn('w:color'), '000000')   # цвет
            tcBorders.append(border)

        # Добавляем в ячейку
        tcPr.append(tcBorders)


    def set_vertical_alignment_to_cell(self, cell: _Cell, align: str = 'center'):
        """
        Устанавливает вертикальное выравнивание ячейки.

        Args:
            cell: объект ячейки python-docx (например, table.cell(0, 0))
            align: 'top', 'center', 'bottom'
        """
        # Допустимые значения
        valid_align = {'top', 'center', 'bottom'}
        if align not in valid_align:
            raise ValueError(f"align must be one of {valid_align}, got '{align}'")

        # Получаем tc (ячейку) и её свойства
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()  # Получаем или создаём <w:tcPr>

        # Удаляем старое выравнивание, если есть
        v_align = tcPr.first_child_found_in("w:vAlign")
        if v_align is not None:
            tcPr.remove(v_align)

        # Создаём новый элемент <w:vAlign w:val="center"/>
        v_align = OxmlElement('w:vAlign')
        v_align.set(qn('w:val'), align)  # например, w:val="center"
        tcPr.append(v_align)


    def delete_row_in_table(self, table: Table, index: int):
        """Удаляет строку из таблицы по указанному индексу.

        Args:
            table (Table): Таблица, из которой нужно удалить строку.
            index (int): Индекс строки, которую нужно удалить.
                        Поддерживает отрицательные индексы (например, -1 — последняя строка).

        Raises:
            IndexError: Если индекс вне допустимого диапазона.
            ValueError: Если таблица пустая.
        """
        tbl = table._tbl  # Внутренний XML-объект таблицы (CT_Tbl)
        rows = tbl.tr_lst  # Список всех строк (lxml elements)

        if not rows:
            raise ValueError("Нельзя удалить строку из пустой таблицы.")

        # Поддержка отрицательных индексов
        if index < 0:
            index = len(rows) + index

        if index < 0 or index >= len(rows):
            raise IndexError(f"Индекс {index} вне допустимого диапазона [0, {len(rows) - 1}]")

        # Получаем XML-элемент строки, которую нужно удалить
        row_to_delete = rows[index]

        # Удаляем строку из XML
        row_to_delete.getparent().remove(row_to_delete)

        # python-docx не обновляет кэш строк автоматически, но table.rows
        # перестраивается при каждом обращении, так что всё будет ок


    def get_paragraph(self, index: int) -> Paragraph:
        """Возвращает параграф по индексу
        """
        return self.doc.paragraphs[index]


    def get_table(self, index: int) -> Table:
        """Возвращает таблицу по индексу
        """
        return self.doc.tables[index]


    def merge_table_cells(self, c1: _Cell, c2: _Cell):
        c1.merge(c2)
        """Удаляет все параграфы из ячейки"""
        """ЭТОТ КОД МОЖЕТ ПРИВЕСТИ К ОШИБКАМ"""
        for paragraph in c1.paragraphs[1:]:
            p = paragraph._element
            p.getparent().remove(p)


    def print_table(self, table: Table):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                print(f"Строка {r}, столбец {c}")

                for paragraph in cell.paragraphs:
                    runs = []
                    for run in paragraph.runs:
                        runs.append(run.text)

                    print(f"runs: {runs}")
            print()


    def insert_paragraph_after_table(self, src_table: Table) -> Paragraph:
        new_paragraph = self.doc.add_paragraph()

        src_table._element.addnext(new_paragraph._element)

        return new_paragraph


    def insert_table_after_table(self, src_table: Table) -> Table:
        new_table = self.doc.add_table(rows=len(src_table.rows), cols=len(src_table.row_cells(0)))
        src_table._element.addnext(new_table._element)
        return new_table


    def clone_table(self, source_table: Table, target_table: Table):
        # Копируем стиль таблицы
        target_table.style = source_table.style
        for source_row, target_row in zip(source_table.rows, target_table.rows):
            self.clone_table_row(source_row, target_row)


"""
    def copy_table(self, table: Table):
        # Создает копию таблицы со всем содержимым
        # Создаем новую таблицу с тем же количеством строк и столбцов
        new_table = table._parent.add_table(rows=0, cols=len(table.columns))

        # Копируем каждую строку
        for row in table.rows:
            new_row = new_table.add_row()
            for i, cell in enumerate(row.cells):
                # Копируем текст
                new_row.cells[i].text = cell.text

                # Копируем форматирование параграфов
                if cell.paragraphs:
                    for j, para in enumerate(cell.paragraphs):
                        if j < len(new_row.cells[i].paragraphs):
                            new_para = new_row.cells[i].paragraphs[j]
                        else:
                            new_para = new_row.cells[i].add_paragraph()

                        # Копируем выравнивание
                        new_para.alignment = para.alignment

                        # Копируем содержимое runs
                        for run in para.runs:
                            new_run = new_para.add_run(run.text)
                            # Копируем форматирование текста
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            new_run.font.name = run.font.name
                            new_run.font.size = run.font.size
                            new_run.font.color.rgb = run.font.color.rgb if run.font.color.rgb else None

        return new_table
"""


    # def copy_text_from_cell(self, cell: _Cell) -> str:
    #     text = ""
    #     for paragraph in cell.paragraph

