import json
from copy import deepcopy
from docx import Document
from docx.shared import Emu, Inches
from docx.table import Table
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
from LegalDocInspector.legal_doc_inspector.doc_creator.docx_editor import DocxRedactor
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table, _Cell, _Row, _Column
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_ROW_HEIGHT_RULE


class CalculationClaimGenerator:
    """
    Класс для создания документа с расчётом для иска
    
    """
    def __init__(self) -> None:
        self.redactor: DocxRedactor = DocxRedactor()
        self.config: dict = None
        self.config2: dict = None
        self.doc: Document = None


    def make_instance(self, config: str, config2: str, template_filename: str, output_filename: str = None):
        """
        config: str - конфиги для шаблона
        config2: str - конфиги для нижней таблицы
        template_filename: str - имя файла-шаблона
        output_filename: str - имя файла, созданного по шаблону
        """
        self.config = self.convert_data_from_calculator(config)
        self.config2 = config2
        # with open("config2.json", "w") as file:
        #     json.dump(config2, file, ensure_ascii=False, indent=4)

        self.redactor.clone_file(template_filename, output_filename)
        self.doc = self.redactor.open(output_filename)

        self.fill_second_table()
        self.fill_other_parts()
        self.fill_file()
        self.redactor.save()
        self.redactor.close()


    def convert_data_from_calculator(self, contracts):
        """На вход метод принимает данные из калькулятора (списком, т.е. в contracts может находится несколько
        наборов данных полученных от калькулятора), каждый элемент списка - это данные об одном контракте.
        
        На выходе возвращаются те же самые данные, но в структуре, удобной для записи в расчёт к иску.
        """

        # with open("temp.json", "w") as file:
        #     json.dump(contracts, file, ensure_ascii=False, indent=4)

        converted_contracts = {"contracts": []}

        for contract in contracts:
            converted_contract = {
                "contract_number": contract["contract_number"],
                "start_date_of_delay": contract["start_of_table"]["start"],
                "end_date_of_delay": contract["start_of_table"]["end"],
                "total_debt": contract["end_of_table1"]["money"],
                "total_peny": contract["end_of_table2"]["money"],
                "periods": [],
            }

            for period in contract.keys():
                if period not in ["contract_number", "start_of_table", "end_of_table1", "end_of_table2", "debt_info"]:
                    new_period = {}
                    new_period["period"] = period

                    for item in contract[period]:

                        if item["text"] == "Итого:":
                            new_period["total"] = item["penalty"]

                    new_period["rows"] = []
                    for item in contract[period]:

                        if "type" in item.keys():
                            if item["type"] == "debt_info":
                                continue

                        if item["text"] != "Итого:":
                            new_period["rows"].append(item)

                    converted_contract["periods"].append(new_period)

            converted_contracts["contracts"].append(converted_contract)

        return converted_contracts

    def _localname(self, elem):
        """Возвращает локальное имя XML-тега (без namespace)."""
        tag = elem.tag
        if '}' in tag:
            return tag.rsplit('}', 1)[1]
        return tag

    def delete_table_and_next_paragraph(self, table):
        """
        Удаляет таблицу и (если существует) _соседний_ параграф, который идёт
        сразу после неё в том же родительском контейнере.
        Не трогает параграфы в других контейнерах и не пытается удалять
        'далёкие' параграфы (только непосредственный сосед).
        """
        tbl_elm = table._element
        parent = tbl_elm.getparent()
        if parent is None:
            return  # нечего удалять

        # Получаем следующий соседний элемент в том же parent
        next_elm = tbl_elm.getnext()

        # Удаляем таблицу
        parent.remove(tbl_elm)

        # Если следующий элемент существует и это параграф (w:p), удаляем его.
        # (Проверяем локальное имя тега — надёжнее, чем endswith)
        if next_elm is not None and self._localname(next_elm) == 'p':
            # Убедимся, что у параграфа тот же родитель (на всякий случай)
            if next_elm.getparent() is parent:
                parent.remove(next_elm)

    def clear_document_from_index(self, index: int, keep_tail=2):
        """
        Удаляет все таблицы, начиная с index, до len(doc.tables) - keep_tail.
        Работает надёжно, т.к. всегда берет таблицу по индексу (живой список).
        - doc: Document
        - index: индекс первой таблицы для удаления (0-based)
        - keep_tail: количество таблиц, которые нужно оставить в конце (по умолчанию 2)
        """
        # Простейшая защита от некорректных значений
        if index < 0:
            raise ValueError("index должен быть >= 0")
        if keep_tail < 0:
            raise ValueError("keep_tail должен быть >= 0")

        # Удаляем в цикле: после удаления следующая таблица займёт тот же индекс.
        # Останавливаемся, когда остаётся только keep_tail таблиц справа.
        while len(self.redactor.doc.tables) > index + keep_tail:
            tbl = self.redactor.doc.tables[index]
            self.delete_table_and_next_paragraph(tbl)

    def fill_file(self):
        # self.clone_table(len(self.config["contracts"]) - 1)

        # self.redactor.print_table(self.redactor.get_table(0))

        for i, contract in enumerate(self.config["contracts"]):
    ### TODO 
            # self._create_table_from_calculation_info_and_replace(contract_info=contract, table_index=i)
            self.fill_contract(contract=contract, contract_index=i)

        self.clear_document_from_index(len(self.config['contracts']))

    def fill_contract(self, contract, contract_index):
        table = self.redactor.get_table(contract_index)
        self._correct_table_height(table)
        # print(self._get_dimensions_of_table(table))
        self.fill_common_contract_info(contract, table)

        self.clone_block(table, 4, len(contract["periods"]) - 1)

        filled_rows = 0     # число заполненных строк во всей таблице
        for i, period in enumerate(contract["periods"]):
            res = self.fill_period(period, i, table, filled_rows)
            filled_rows += res
            # filled_rows += 3
            # payments += len(period["payments_1"])
        #     payments += len(period["payments_2"])
        self._correct_table_width(table, contract_index)

    def fill_common_contract_info(self, contract:dict, table: Table):
        to_replace = (
            ("CONTRACTNUMBER",            contract["contract_number"]),
            ("STARTDELAYDATE",     contract["start_date_of_delay"]),
            ("ENDDELAYDATE",      contract["end_date_of_delay"]),
            (self.borders("сумма долга"),               contract["total_debt"]),
            (self.borders("сумма пеней"),               contract["total_peny"]),
        )

        cell = table.row_cells(0)[3]
        flag = self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])
        cell = table.row_cells(1)[3]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[1][0], to_replace[1][1])
        cell = table.row_cells(1)[10]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[2][0], to_replace[2][1])
        cell = table.row_cells(8)[0]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[3][0], to_replace[3][1])
        cell = table.row_cells(9)[0]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[4][0], to_replace[4][1])


    def clone_block(self, table: Table, index: int, count: int = 1):
        """Клонирует блок. 1 блок - 1 период

        Args:
            index (int): индекс первой строки копируемого блока
            count (int): количество копий.
        """

        stride = 4
        for i in range(count):
            for j in range(stride):
                new_row = self.redactor.insert_row_in_table(table, index + (i + 1) * stride + j)
                self.redactor.clone_table_row(table.rows[index + j], new_row)

            self.merge_row_5(table, index + (i + 1) * stride)
            self.merge_row_6(table, index + (i + 1) * stride + 1)
            self.merge_row_5(table, index + (i + 1) * stride + 2, True)
            self.merge_row_6(table, index + (i + 1) * stride + 3)


    def fill_period(self, period, period_index, table: Table, filled_rows: int):
        self.fill_common_period_info(period, period_index, table, filled_rows)

        """
        template_offset - это индекс, по которому находится строка-шаблон header-строки. Он вычисляется
        следующим образом. 4 - это первые четыре строки в таблице, они содержат информацию о договоре.
        filled_rows - количество строк, которые заполнены в таблице (в это число не входят первые четыре 
        строки в таблице, которые не относятся не к каким периодам)
        """
        template_offset_1 = 4 + filled_rows
        template_offset_2 = template_offset_1 + 1

        # Этот флаг отвечает за то какую строку мы заполняем
        first_row_flag = False

        current_filled_rows = 0
        for row in period["rows"]:
            paste_offset = template_offset_1 + current_filled_rows + 3    # +3 - это плюс три строки с шаблонами
            if first_row_flag:
                paste_offset -= 1
            # print(f"Обрабатываем строку: {row}")

            if self.type_of_row(row) == 1:
                if not first_row_flag:
                    self.fill_row_type_1(table, template_offset_1, row)
                    first_row_flag = True
                else:
                    self.clone_row_type_1(table, template_offset_1 + 2, paste_offset)
                    self.fill_row_type_1(table, paste_offset, row)
                    # template_offset_1 = paste_offset

            elif self.type_of_row(row) == 2:
                self.clone_row_type_2(table, template_offset_1 + 1, paste_offset)
                self.fill_row_type_2(table, paste_offset, row)
                # template_offset_2 = paste_offset

            current_filled_rows += 1
            # paste_offset += 1
            # self.redactor.save()

        # Здесь мы добавляем единицу, поскольку каждый блок заканчивается строкой "итого"
        current_filled_rows += 1
        # Удаляем строки с шаблонами
        self.redactor.delete_row_in_table(table, template_offset_1 + 1)
        self.redactor.delete_row_in_table(table, template_offset_1 + 1)
        # current_filled_rows += 2

        return current_filled_rows


    def fill_common_period_info(self, period, period_index: int, table: Table, filled_rows: int):
        # print(f"месяц: {period['period']}")
        to_replace = (
            ("MONTHPERIOD",              period["period"]),
            (self.borders("сумма пени за период"),          period["total"]),
        )

        cell = table.row_cells(4 + filled_rows)[0]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])

        cell = table.row_cells(4 + filled_rows + 3)[11]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[1][0], to_replace[1][1])


    def type_of_row(self, row: dict):
        """Возвращает тип строки. Если в строке есть текст, значит это строка-заголовок и возвращаем 1.
        Если текста нет, значит это строка в которой есть формула, ставка, пени и тд, возвращаем 2
        1 - строка-заголовок
        2 - строка-куча чисел

        Args:
            row (dict): строка row представлена в виде словаря с полями, в том числе есть поле "text".

        Returns:
            _type_: 1 или 2
        """
        if row["text"] is None:
            return 2
        else:
            return 1


    def count_row_by_types(self, rows):
        n_header_rows = 0
        n_number_rows = 0
        for row in rows:
            row_type = self.type_of_row(row)
            if row_type == 1:
                n_header_rows += 1
            elif row_type == 2:
                n_number_rows += 1
            else:
                print(f"Строка row={row} не соответствует ни одному формату, не ясно как её обрабатывать.")
                raise RuntimeError(f"Invalid row: {row}")
        return (n_header_rows, n_number_rows)


    def clone_row_type_1(self, table: Table, source_row_index: int, _target_row_index: int = None, merge_flag: bool = False):
        """Клонирует хедер-строку.

        Args:
            table (Table): Таблица, в которой происходят изменения
            source_row_index (int, optional): Индекс строки, которую копируем (она является шаблонной).
            target_row_index (int, optional): Индекс, по которому вставляем копию строки. По 
            умолчанию вставляется сразу после шаблонной строки.
            
            Update
            merge_flag - нужно ли присоединять строку к первому столбцу
        """
        target_row_index = source_row_index if _target_row_index is None else _target_row_index
        new_row = self.redactor.insert_row_in_table(table, target_row_index)
        self.redactor.clone_table_row(table.rows[source_row_index], new_row)
        self.merge_row_5(table, target_row_index, merge_flag)
        self.redactor.merge_table_cells(table.rows[target_row_index - 1].cells[0], table.rows[target_row_index].cells[0])


    def clone_row_type_2(self, table: Table, source_row_index: int, _target_row_index: int = None):
        """Клонирует строку с кучей численных данных.

        Args:
            table (Table): Таблица, в которой происходят изменения
            source_row_index (int, optional): Индекс строки, которую копируем (она является шаблонной).
            target_row_index (int, optional): Индекс, по которому вставляем копию строки. По 
            умолчанию вставляется сразу после шаблонной строки.
        """
        target_row_index = source_row_index if _target_row_index is None else _target_row_index
        new_row = self.redactor.insert_row_in_table(table, target_row_index)
        self.redactor.clone_table_row(table.rows[source_row_index], new_row)
        self.merge_row_6(table, target_row_index)
        self.redactor.merge_table_cells(table.rows[target_row_index - 1].cells[0], table.rows[target_row_index].cells[0])


    def fill_row_type_1(self, table: Table, row_index: int, row: dict):
        to_replace = (
            (self.borders("сумма начисления"),              row["debt"]),
            (self.borders("дата начала начисления"),        row["period"][0]),
            (self.borders("вставляемый текст"),             row["text"]),
        )

        cell = table.row_cells(row_index)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])
        cell = table.row_cells(row_index)[3]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[1][0], to_replace[1][1])
        cell = table.row_cells(row_index)[4]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[2][0], to_replace[2][1])


    def fill_row_type_2(self, table: Table, row_index: int, row: dict):
        # print(f"row={row}")
        to_replace = (
            (self.borders("сумма начисления"),              row["debt"]),
            (self.borders("дата начала начисления"),        row["period"][0]),
            (self.borders("дата конца начисления"),         row["period"][1]),
            (self.borders("кол-во дней начисления"),        str(row["period"][2])),
            (self.borders("ставка"),                        row["penalty_period_info"][0]),
            (self.borders("доля ставки"),                   row["penalty_period_info"][1]),
            (self.borders("формула"),                       row["formulae"]),
            (self.borders("пени"),                          row["penalty"]),
        )

        cell = table.row_cells(row_index)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])
        cell = table.row_cells(row_index)[3]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[1][0], to_replace[1][1])
        cell = table.row_cells(row_index)[4]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[2][0], to_replace[2][1])
        cell = table.row_cells(row_index)[5]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[3][0], to_replace[3][1])

        cell = table.row_cells(row_index)[6]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[4][0], to_replace[4][1])
        cell = table.row_cells(row_index)[8]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[5][0], to_replace[5][1])
        cell = table.row_cells(row_index)[9]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[6][0], to_replace[6][1])
        cell = table.row_cells(row_index)[11]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[7][0], to_replace[7][1])


    def borders(self, string: str, start: str = "/*", end: str = "*/"):
        return start + string + end


    def clone_table(self, count: int = 1):
        # Берём первую таблицу
        original_table = self.redactor.get_table(0)

        for i in range(count):
            # Получаем XML-элемент таблицы
            table_element = original_table._element

            # Глубокая копия XML-элемента таблицы (включая все внутренние элементы)
            cloned_table_element = deepcopy(table_element)

            # Находим родительский элемент (обычно это <w:body>)
            parent_element = table_element.getparent()

            # Определяем позицию исходной таблицы
            index_in_parent = parent_element.index(table_element)

            # Создаём пустой параграф
            new_paragraph = self.doc.add_paragraph()
            new_paragraph_element = new_paragraph._element

            # Создаём клонированную таблицу (обёртка для XML-элемента)
            # from docx.table import Table
            cloned_table = Table(cloned_table_element, self.doc)

            # Вставляем пустой параграф сразу после исходной таблицы
            parent_element.insert(index_in_parent + 1, new_paragraph_element)

            # Вставляем клонированную таблицу после параграфа
            parent_element.insert(index_in_parent + 2, cloned_table_element)


    def merge_row_1(self, table, row_index):
        # Объединяем ячейки первой строки
        row = table.rows[row_index]
        self.redactor.merge_table_cells(row.cells[0], row.cells[1])
        for i in range(3, 12):
            self.redactor.merge_table_cells(row.cells[2], row.cells[i])


    def merge_row_2(self, table, row_index):
        # Объединяем ячейки второй строки
        row = table.rows[row_index]
        self.redactor.merge_table_cells(row.cells[0], row.cells[1])
        for i in range(3, 7):
            self.redactor.merge_table_cells(row.cells[2], row.cells[i])
        for i in range(8, 10):
            self.redactor.merge_table_cells(row.cells[7], row.cells[i])
        self.redactor.merge_table_cells(row.cells[10], row.cells[11])


    def merge_row_3_4(self, table, row_3_index):
        # Объединяем ячейки третьей и четвертой строк
        row_3 = table.rows[row_3_index]
        row_4 = table.rows[row_3_index + 1]
        for row in [row_3, row_4]:
            self.redactor.merge_table_cells(row.cells[1], row.cells[2])
            self.redactor.merge_table_cells(row.cells[6], row.cells[7])
            self.redactor.merge_table_cells(row.cells[9], row.cells[10])

        for i in range(4, 6):
            self.redactor.merge_table_cells(row_3.cells[3], row_3.cells[i])

        for i in [0, 1, 6, 7, 8, 9, 10, 11]:
            self.redactor.merge_table_cells(row_3.cells[i], row_4.cells[i])


    def merge_row_5(self, table, row_index, first_cell_flag = False):
        """
        first_cell_flag - флаг сообщает, нужно ли присоединять первую колонку текущей строки
        к общей колонке "период". По умолчанию не нужно.
        """
        row = table.rows[row_index]
        self.redactor.merge_table_cells(row.cells[1], row.cells[2])
        for i in range(5, 12):
            self.redactor.merge_table_cells(row.cells[4], row.cells[i])

        if first_cell_flag:
            previous_row = table.rows[row_index - 1]
            self.redactor.merge_table_cells(previous_row.cells[0], row.cells[0])


    def merge_row_6(self, table, row_index):
        row = table.rows[row_index]
        self.redactor.merge_table_cells(row.cells[1], row.cells[2])
        self.redactor.merge_table_cells(row.cells[6], row.cells[7])
        self.redactor.merge_table_cells(row.cells[9], row.cells[10])

        previous_row = table.rows[row_index - 1]
        self.redactor.merge_table_cells(previous_row.cells[0], row.cells[0])


    def merge_row_8(self, table, row_index):
        row = table.rows[row_index]
        for i in range(1, 12):
            self.redactor.merge_table_cells(row.cells[0], row.cells[i])


    def clone_payment_row(self, table: Table, count: int = 1, source_row_index: int = 5, target_row_index: int = None):
        """Клонирует строку с данными платежа.

        Args:
            table (Table): Таблица, в которой происходят изменения
            source_row_index (int, optional): Индекс строки, которую копируем (она является шаблонной). 
            По умолчанию = 5.
            target_row_index (int, optional): Индекс, по которому вставляем копию строки. По 
            умолчанию вставляется сразу после шаблонной строки.
        """

        _target_row_index = source_row_index + 1 if target_row_index is None else target_row_index

        for i in range(count):
            new_row = self.redactor.insert_row_in_table(table, _target_row_index)
            self.redactor.clone_table_row(table.rows[source_row_index], new_row)
            self.merge_row_6(table, _target_row_index)
            self.redactor.merge_table_cells(table.rows[_target_row_index - 1].cells[0], table.rows[_target_row_index].cells[0])


    def fill_second_table(self):
        self.fill_second_table_common_info()

        table = self.redactor.get_table(-2)
        # self.redactor.print_table(table)
        
        # по сути число договоров
        n_contracts = len(self.config2["contracts_info"])

        # По смыслу, это количство уже заполненных строк в таблице. Начинаем не с 0, а с 1 потому что
        # в таблице есть заголовочная строка, в которой ничего заполнять не нужна, но она есть сама 
        # по себе и её нужно учитывать. По сути, значение n_row равно индексу строки, которую мы будем 
        # заполнять на текущей итерации
        n_rows = 1
        for i, contract in enumerate(self.config2["contracts_info"]):

            if i < n_contracts - 1:
                self.second_table_clone_row(table, n_rows)

            # contract_number = list(self.config2["table_info"].keys())[i]
            contract_number = self.config2["contracts_info"][i][1]

            if self.config2["table_info"][contract_number]["correcting_debt"] == "0,00":
                self.fill_second_table_simple_row(table, n_rows, i)
                n_rows += 1
            
            elif self.config2["table_info"][contract_number]["accrual_debt"] == "0,00":
                self.fill_second_table_row_with_year_adjustment_only(table, n_rows, i)
                n_rows += 1

            else:
                self.second_table_clone_row(table, n_rows)
                self.fill_second_table_complex_row(table, n_rows, i)
                n_rows += 2


    def fill_second_table_common_info(self):
        table = self.redactor.get_table(-2)
        self.redactor.replace_text_in_paragraph(
            table.row_cells(2)[2].paragraphs[0],
            self.borders("сумма долга"),
            self.config2["table_info"]["all_debt"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(2)[3].paragraphs[0],
            self.borders("неустойка общая"),
            self.config2["table_info"]["all_penalty"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(2)[4].paragraphs[0],
            self.borders("цена иска"),
            self.config2["table_info"]["cost_of_lawsuit"]
        )


    def second_table_clone_row(self, table: Table, cloning_row_index: int):
        new_row = self.redactor.insert_row_in_table(table, cloning_row_index + 1)
        self.redactor.clone_table_row(table.rows[cloning_row_index], new_row)


    def fill_second_table_simple_row(self, table: Table, row_index: int, contract_index: int):
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[0].paragraphs[0],
            self.borders("номер договора"),
            self.config2["contracts_info"][contract_index][1]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[1].paragraphs[0],
            self.borders("период"),
            self.config2["contracts_info"][contract_index][2]["contract_periods"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config2["contracts_info"][contract_index][2]["debt"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[3].paragraphs[0],
            self.borders("неустойка"),
            # "#неустойка#".upper()
            self.config2["contracts_info"][contract_index][2]["penalty"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[4].paragraphs[0],
            self.borders("неустойка+задолженность"),
            # "#неустойка+задолженность#".upper()
            self.config2["contracts_info"][contract_index][2]["debt_penalty"]
        )


    def fill_second_table_complex_row(self, table: Table, row_index: int, contract_index: int):
        self.second_table_merge_rows(table, row_index)

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[0].paragraphs[0],
            self.borders("номер договора"),
            self.config2["contracts_info"][contract_index][1]
        )

        # Ячейка Период
        source_paragraph = table.row_cells(row_index)[1].paragraphs[0]
        new_paragraph = table.row_cells(row_index)[1].add_paragraph("")
        self.redactor.clone_paragraph(source_paragraph, new_paragraph)
        self.redactor.replace_text_in_paragraph(
            source_paragraph,
            self.borders("период"),
            self.config2["contracts_info"][contract_index][2]["contract_periods"]
        )
        self.redactor.replace_text_in_paragraph(
            new_paragraph,
            self.borders("период"),
            "текущие начисления"
        )
        self.redactor.paragraph_text_set_bold(new_paragraph)

        source_paragraph = table.row_cells(row_index + 1)[1].paragraphs[0]
        new_paragraph = table.row_cells(row_index + 1)[1].add_paragraph("")
        self.redactor.clone_paragraph(source_paragraph, new_paragraph)
        self.redactor.replace_text_in_paragraph(
            source_paragraph,
            self.borders("период"),
            self.config2["contracts_info"][contract_index][2]["contract_periods_correcting"]
        )
        self.redactor.replace_text_in_paragraph(
            new_paragraph,
            self.borders("период"),
            "доля от ГК"
        )
        self.redactor.paragraph_text_set_bold(new_paragraph)

        # Ячейка Задолженность
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config2["contracts_info"][contract_index][2]["accrual_debt"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index + 1)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config2["contracts_info"][contract_index][2]["correcting_debt"]
        )

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[3].paragraphs[0],
            self.borders("неустойка"),
            # "#неустойка#".upper()
            self.config2["contracts_info"][contract_index][2]["penalty"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[4].paragraphs[0],
            self.borders("неустойка+задолженность"),
            # "#неустойка+задолженность#".upper()
            self.config2["contracts_info"][contract_index][2]["debt_penalty"]
        )


    def fill_second_table_row_with_year_adjustment_only(self, table: Table, row_index: int, contract_index: int):
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[0].paragraphs[0],
            self.borders("номер договора"),
            self.config2["contracts_info"][contract_index][1]
        )

        source_paragraph = table.row_cells(row_index)[1].paragraphs[0]
        new_paragraph = table.row_cells(row_index)[1].add_paragraph("")
        self.redactor.clone_paragraph(source_paragraph, new_paragraph)
        self.redactor.replace_text_in_paragraph(
            source_paragraph,
            self.borders("период"),
            self.config2["contracts_info"][contract_index][2]["contract_periods_correcting"]
        )
        self.redactor.replace_text_in_paragraph(
            new_paragraph,
            self.borders("период"),
            "доля от ГК"
        )
        self.redactor.paragraph_text_set_bold(new_paragraph)

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config2["contracts_info"][contract_index][2]["correcting_debt"]
        )

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[3].paragraphs[0],
            self.borders("неустойка"),
            # "#неустойка#".upper()
            self.config2["contracts_info"][contract_index][2]["penalty"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[4].paragraphs[0],
            self.borders("неустойка+задолженность"),
            # "#неустойка+задолженность#".upper()
            self.config2["contracts_info"][contract_index][2]["debt_penalty"]
        )


    def second_table_merge_rows(self, table: Table, row_index: int):
        row_1 = table.rows[row_index]
        row_2 = table.rows[row_index + 1]
        for i in [0, 3, 4]:
            self.redactor.merge_table_cells(row_1.cells[i], row_2.cells[i])


    def fill_other_parts(self):
        to_replace = {
            "/*цена иска*/": self.config2["lawsuit_info"]["cost"],
            "/*госпошлина*/": self.config2["lawsuit_info"]["tax"],
            "/*текущая дата*/": datetime.strptime(self.config2["current_date"], "%Y-%m-%d").strftime("%d.%m.%Y")
        }

        for paragraph in self.doc.paragraphs:
            # runs = []
            for run in paragraph.runs:
                # runs.append(run.text)
                for k, v in to_replace.items():
                    if k in run.text:
                        run.text = run.text.replace(k, v)
            # print(f"runs={runs}")


    def _get_dimensions_of_table(self, table:Table):
        """Получает все размеры таблицы"""
        
        # Размеры таблицы
        dimensions = {
            'rows_count': len(table.rows),
            'columns_count': len(table.columns),
            'column_widths': [],
            'row_heights': [],
            'cell_dimensions': []
        }
        
        # Получение ширины колонок
        for col_idx, column in enumerate(table.columns):
            # Ширина колонки (может быть None если не задана явно)
            width = column.width
            dimensions['column_widths'].append({
                'column_index': col_idx,
                'width': width,
                'width_in_inches': width.inches if width else None
            })
        
        # Получение высоты строк
        for row_idx, row in enumerate(table.rows):
            height = row.height
            dimensions['row_heights'].append({
                'row_index': row_idx,
                'height': height,
                'height_in_inches': height.inches if height else None
            })
        
        # Получение размеров каждой ячейки
        for row_idx, row in enumerate(table.rows):
            row_cells = []
            for col_idx, cell in enumerate(row.cells):
                # Информация о ячейке
                cell_info = {
                    'row': row_idx,
                    'column': col_idx,
                    'text': cell.text.strip(),
                    'colspan': 1,  # по умолчанию
                    'rowspan': 1   # по умолчанию
                }
                row_cells.append(cell_info)
            dimensions['cell_dimensions'].append(row_cells)
        
        return dimensions
    
    def _correct_table_height(self, table:Table):
        for row in table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Emu(205570)

    def _correct_table_width(self, table:Table, table_number:int): 
        
        # if table_number==0:
        #     self.set_exact_cell_dimensions(table.rows[0].cells[0], Inches(2.5).twips)
        #     self.set_exact_cell_dimensions(table.rows[2].cells[0], Inches(1.48).twips) 
        #     self.set_exact_cell_dimensions(table.rows[2].cells[1], Inches(1.4).twips)
        #     self.set_exact_cell_dimensions(table.rows[3].cells[3], Inches(1.5).twips)
        #     self.set_exact_cell_dimensions(table.rows[3].cells[4], Inches(1.5).twips)
        #     self.set_exact_cell_dimensions(table.rows[3].cells[5], Inches(0.8).twips)
        #     self.set_exact_cell_dimensions(table.rows[3].cells[6], Inches(0.5).twips)
        #     self.set_exact_cell_dimensions(table.rows[3].cells[7], Inches(0.5).twips)
        #     self.set_exact_cell_dimensions(table.rows[3].cells[8], Inches(2).twips)
        #     self.set_exact_cell_dimensions(table.rows[3].cells[11], Inches(1.8).twips) 
        # else:

        # for i, cell in enumerate(table.rows[2].cells):
        #     print(i, cell.text)
        # for i, cell in enumerate(table.rows[3].cells):
        #     print(i, cell.text)
        
        self.set_exact_cell_dimensions(table.rows[0].cells[0], Inches(4.2).twips) #начало просрочки

        self.set_exact_cell_dimensions(table.rows[2].cells[0], Inches(2.5).twips) #месяц
        self.set_exact_cell_dimensions(table.rows[2].cells[1], Inches(2).twips) #долг
        self.set_exact_cell_dimensions(table.rows[2].cells[1], Inches(2.5).twips) #долг

        self.set_exact_cell_dimensions(table.rows[2].cells[9], Inches(3).twips) #долг
        self.set_exact_cell_dimensions(table.rows[2].cells[6], Inches(1.4).twips) #долг
        self.set_exact_cell_dimensions(table.rows[2].cells[8], Inches(1.4).twips) #долг
        self.set_exact_cell_dimensions(table.rows[2].cells[11], Inches(2.5).twips) #долг
        self.set_exact_cell_dimensions(table.rows[2].cells[9], Inches(6).twips) #долг

        self.set_exact_cell_dimensions(table.rows[3].cells[3], Inches(2).twips) # c
        self.set_exact_cell_dimensions(table.rows[3].cells[4], Inches(2).twips) # по
        self.set_exact_cell_dimensions(table.rows[3].cells[5], Inches(1.2).twips) # дней
        # self.set_exact_cell_dimensions(table.rows[3].cells[6], Inches(0.5).twips) # ставка
        # self.set_exact_cell_dimensions(table.rows[3].cells[8], Inches(2).twips) # доля
        # self.set_exact_cell_dimensions(table.rows[2].cells[11], Inches(3.4).twips)  # пени
        
        # print('-----------------')

    def set_exact_cell_dimensions(self, cell, width=None, height=None):
        """Устанавливает точные размеры ячейки"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        if width:
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(width))
            tcW.set(qn('w:type'), 'dxa') # в TWIPS (1/20 point) 
            tcPr.append(tcW) 
        if height: # Для высоты нужно работать со строкой 
            pass



    
    def _create_table_from_calculation_info_and_replace(self, contract_info:dict, table_index):
        old_table = self.redactor.get_table(table_index)
        doc = self.redactor.doc
        parent = old_table._element.getparent()
        index = parent.index(old_table._element)
        new_table = doc.add_table(rows=1, cols=13, style = old_table.style )
        parent.remove(old_table._element)
        self._create_table_title(new_table, contract_info)
        
        # Вставляем на место
        parent.insert(index, new_table._element)

        # self._create_table_month_info()
            

    def _create_table_title(self, table:Table, contract_info:dict):
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[0].height = Cm(0.6)
        row_cells = table.rows[0].cells
        row_cells[0].merge(row_cells[2])
        self._put_text_into_table_cell('Информация о расчёте',
                                        row_cells[0],
                                        need_gray_bgc=True,
                                        orient='left')
        
    






    def _put_text_into_table_cell(self, text:str, cell:_Cell, font_size=9, need_bold=False, need_italic=False, orient="center", need_vertical_orient=True, need_gray_bgc=False):
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = need_bold
        run.italic = need_italic
        run.font.name = 'Times New Roman'

        element = run._element
        rPr = element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        # run._element.rPr.rFonts.set('w:eastAsia', 'Times New Roman')  # для корректного отображения в Word
        run.font.size = Pt(font_size)

        match orient:
            case "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            case "left":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            case "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if need_vertical_orient:
            self._set_cell_vertical_alignment(cell=cell)

        if need_gray_bgc:
            self._set_cell_background(cell=cell)


        

    def _set_cell_vertical_alignment(self, cell:_Cell, align="center"):
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tag = tc_pr.xpath('w:vAlign')
        if tag:
            el = tag[0]
            el.set(qn('w:val'), align)
        else:
            valign = OxmlElement('w:vAlign')
            valign.set(qn('w:val'), align)
            tc_pr.append(valign)

    
    def _set_cell_background(self, cell:_Cell, color_hex = "#d5d5d5"):
        tc_pr = cell._tc.get_or_add_tcPr()
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        tc_pr.append(shading_elm)

    def _set_row_bottom_border(self, row:_Row, color="000000", size=4):
        """
        :param row: объект строки (Row)
        :param color: цвет в HEX без решётки (#), например "000000" — чёрный
        :param size: толщина линии (в 1/8 pt, например 4 = 0.5pt, 6 = 0.75pt)
        """
        tr = row._tr
        tc_borders = OxmlElement('w:tblBottomBorders')

        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), str(size))
        bottom_border.set(qn('w:color'), color)

        tc_borders.append(bottom_border)

        for el in tr.xpath('w:tblBottomBorders'):
            tr.remove(el)

        tr.append(tc_borders)
    
