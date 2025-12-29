import json
from datetime import datetime, timedelta
import re

from docx import Document
from docx.table import Table

from LegalDocInspector.legal_doc_inspector.doc_creator.docx_editor import DocxRedactor



class ClaimGenerator:
    """
    Класс создания документа по шаблону. При создании экземпляра принимает три аргумента:
    
    """

    def __init__(self) -> None:
        self.redactor: DocxRedactor = DocxRedactor()
        self.config: dict = None
        self.doc: Document = None


    def make_instance(self, config: str, template_filename: str, output_filename: str = None):
        """
        config_filename: str - конфиги для шаблона
        template_filename: str - имя файла-шаблона
        output_filename: str - имя файла, созданного по шаблону
        """
        self.config = config

        self.redactor.clone_file(template_filename, output_filename)
        self.doc = self.redactor.open(output_filename)
        # with open("claim_config.json", "w") as file:
        #     json.dump(self.config, file, ensure_ascii=False, indent=4)

        self.fill_file()

        self.redactor.save()
        self.redactor.close()


    def parse_config(self, filename: str) -> dict:
        with open(filename, 'r', encoding="utf-8") as file:
            data = json.load(file)
        return data


    def borders(self, string: str, start: str = "/*", end: str = "*/"):
        return start + string + end


    def fill_file(self):
        # self.redactor.print_table(self.redactor.get_table(0))
        self.prepare_data()

        # with open("claim_config.json", "w") as file:
        #     json.dump(self.config, file, ensure_ascii=False, indent=4)

        self.fix_quotes()
        self.fill_first_table()
        self.fill_second_table()
        self.fill_third_table()
        self.fill_first_list()
        self.fill_second_list()
        self.fill_third_list()
        self.fill_other_parts()


    def fill_first_table(self):
        # self.redactor.print_table(table)

        to_replace = (
            # строка 2
            (self.borders("истец полное имя"),      self.config["plaintiff_info"]["full_name"]),
            (self.borders("истец огрн"),            self.config["plaintiff_info"]["ogrn"]),
            (self.borders("истец инн"),             self.config["plaintiff_info"]["inn"]),

            # Строка 3
            (self.borders("истец адрес"),           self.config["plaintiff_info"]["addres"]),

            # Строка 6
            (self.borders("ответчик полное имя"),   self.config["defendant_info"]["full_name"].upper()),
            (self.borders("ответчик огрн"),         self.config["defendant_info"]["ogrn"]),
            (self.borders("ответчик инн"),          self.config["defendant_info"]["inn"]),

            # Строка 7
            (self.borders("ответчик адрес"),        self.config["defendant_info"]["addres"]),
        )
        print(self.config["defendant_info"]["full_name"])
        table = self.redactor.get_table(0)

        cell = table.row_cells(2)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])
        self.redactor.replace_text_in_paragraph(cell.paragraphs[1], to_replace[1][0], to_replace[1][1])
        self.redactor.replace_text_in_paragraph(cell.paragraphs[1], to_replace[2][0], to_replace[2][1])

        cell = table.row_cells(3)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[3][0], to_replace[3][1])

        cell = table.row_cells(6)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[4][0], to_replace[4][1])
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[5][0], to_replace[5][1])
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[6][0], to_replace[6][1])

        cell = table.row_cells(7)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[7][0], to_replace[7][1])

        self.redactor.save()


    def fill_second_table(self):

        table = self.redactor.get_table(1)
        self.second_table_fill_common_info(table)
        # self.redactor.print_table(table)
        
        # по сути число договоров
        n_contracts = len(self.config["contracts_info"])

        # По смыслу, это количство уже заполненных строк в таблице. Начинаем не с 0, а с 1 потому что
        # в таблице есть заголовочная строка, в которой ничего заполнять не нужна, но она есть сама 
        # по себе и её нужно учитывать. По сути, значение n_row равно индексу строки, которую мы будем 
        # заполнять на текущей итерации
        n_rows = 1
        for i, contract in enumerate(self.config["contracts_info"]):

            if i < n_contracts - 1:
                self.third_table_clone_row(table, n_rows)

            # contract_number = list(self.config["table_info"].keys())[i]
            contract_number = self.config["contracts_info"][i][1]

            if self.config["table_info"][contract_number]["correcting_debt"] == "0,00":
                self.second_table_fill_simple_row(table, n_rows, i)
                n_rows += 1

            elif self.config["table_info"][contract_number]["accrual_debt"] == "0,00":
                self.second_table_fill_row_with_year_adjustment_only(table, n_rows, i)
                n_rows += 1

            else:
                self.third_table_clone_row(table, n_rows)
                self.second_table_fill_complex_row(table, n_rows, i)
                n_rows += 2


    def second_table_fill_common_info(self, table: Table):
        self.redactor.replace_text_in_paragraph(
            table.row_cells(2)[2].paragraphs[0],
            self.borders("сумма долга"),
            self.config["table_info"]["all_debt"]
        )


    def second_table_fill_simple_row(self, table: Table, row_index: int, contract_index: int):
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[0].paragraphs[0],
            self.borders("номер договора"),
            self.config["contracts_info"][contract_index][1]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[1].paragraphs[0],
            self.borders("период"),
            self.config["contracts_info"][contract_index][2]["contract_periods"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config["contracts_info"][contract_index][2]["accrual_debt"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[3].paragraphs[0],
            self.borders("срок оплаты"),
            self.config["contracts_info"][contract_index][2]["last_day"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[4].paragraphs[0],
            self.borders("пункт"),
            self.config["contracts_info"][contract_index][2]["contract_point"]
        )


    def second_table_fill_complex_row(self, table: Table, row_index: int, contract_index: int):
        self.third_table_merge_rows(table, row_index)

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[0].paragraphs[0],
            self.borders("номер договора"),
            self.config["contracts_info"][contract_index][1]
        )

        # Ячейка Период
        source_paragraph = table.row_cells(row_index)[1].paragraphs[0]
        new_paragraph = table.row_cells(row_index)[1].add_paragraph("")
        self.redactor.clone_paragraph(source_paragraph, new_paragraph)
        self.redactor.replace_text_in_paragraph(
            source_paragraph,
            self.borders("период"),
            self.config["contracts_info"][contract_index][2]["contract_periods"]
        )
        self.redactor.replace_text_in_paragraph(
            new_paragraph,
            self.borders("период"),
            "текущие начисления"
        )
        self.redactor.paragraph_text_set_bold(new_paragraph)

        source_paragraph = table.row_cells(row_index + 1)[1].paragraphs[0]
        new_paragraph = table.row_cells(row_index + 1)[1].add_paragraph("")
        self.redactor.clone_paragraph(source_paragraph, new_paragraph)
        self.redactor.replace_text_in_paragraph(
            source_paragraph,
            self.borders("период"),
            self.config["contracts_info"][contract_index][2]["contract_periods_correcting"]
        )
        self.redactor.replace_text_in_paragraph(
            new_paragraph,
            self.borders("период"),
            "доля от ГК"
        )
        self.redactor.paragraph_text_set_bold(new_paragraph)

        # Ячейка Задолженность
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config["contracts_info"][contract_index][2]["accrual_debt"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index + 1)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config["contracts_info"][contract_index][2]["correcting_debt"]
        )

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[3].paragraphs[0],
            self.borders("срок оплаты"),
            self.config["contracts_info"][contract_index][2]["last_day"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[4].paragraphs[0],
            self.borders("пункт"),
            self.config["contracts_info"][contract_index][2]["contract_point"]
        )


    def second_table_fill_row_with_year_adjustment_only(self, table: Table, row_index: int, contract_index: int):
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[0].paragraphs[0],
            self.borders("номер договора"),
            self.config["contracts_info"][contract_index][1]
        )

        source_paragraph = table.row_cells(row_index)[1].paragraphs[0]
        new_paragraph = table.row_cells(row_index)[1].add_paragraph("")
        self.redactor.clone_paragraph(source_paragraph, new_paragraph)
        self.redactor.replace_text_in_paragraph(
            source_paragraph,
            self.borders("период"),
            self.config["contracts_info"][contract_index][2]["contract_periods_correcting"]
        )
        self.redactor.replace_text_in_paragraph(
            new_paragraph,
            self.borders("период"),
            "доля от ГК"
        )
        self.redactor.paragraph_text_set_bold(new_paragraph)

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config["contracts_info"][contract_index][2]["correcting_debt"]
        )

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[3].paragraphs[0],
            self.borders("срок оплаты"),
            self.config["contracts_info"][contract_index][2]["last_day"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[4].paragraphs[0],
            self.borders("пункт"),
            self.config["contracts_info"][contract_index][2]["contract_point"]
        )


    def fill_third_table(self):

        table = self.redactor.get_table(2)
        self.third_table_fill_common_info(table)
        # self.redactor.print_table(table)
        
        # по сути число договоров
        n_contracts = len(self.config["contracts_info"])

        # По смыслу, это количство уже заполненных строк в таблице. Начинаем не с 0, а с 1 потому что
        # в таблице есть заголовочная строка, в которой ничего заполнять не нужна, но она есть сама 
        # по себе и её нужно учитывать. По сути, значение n_row равно индексу строки, которую мы будем 
        # заполнять на текущей итерации
        n_rows = 1
        for i, contract in enumerate(self.config["contracts_info"]):

            if i < n_contracts - 1:
                self.third_table_clone_row(table, n_rows)

            # contract_number = list(self.config["table_info"].keys())[i]
            contract_number = self.config["contracts_info"][i][1]

            if self.config["table_info"][contract_number]["correcting_debt"] == "0,00":
                self.third_table_fill_simple_row(table, n_rows, i)
                n_rows += 1

            elif self.config["table_info"][contract_number]["accrual_debt"] == "0,00":
                self.third_table_fill_row_with_year_adjustment_only(table, n_rows, i)
                n_rows += 1

            else:
                self.third_table_clone_row(table, n_rows)
                self.third_table_fill_complex_row(table, n_rows, i)
                n_rows += 2


    def third_table_fill_common_info(self, table: Table):
        # table = self.redactor.get_table(1)
        self.redactor.replace_text_in_paragraph(
            table.row_cells(2)[2].paragraphs[0],
            self.borders("сумма долга"),
            self.config["table_info"]["all_debt"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(2)[3].paragraphs[0],
            self.borders("неустойка общая"),
            self.config["table_info"]["all_penalty"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(2)[4].paragraphs[0],
            self.borders("цена иска"),
            self.config["table_info"]["cost_of_lawsuit"]
        )


    def third_table_clone_row(self, table: Table, cloning_row_index: int):
        new_row = self.redactor.insert_row_in_table(table, cloning_row_index + 1)
        self.redactor.clone_table_row(table.rows[cloning_row_index], new_row)


    def third_table_fill_simple_row(self, table: Table, row_index: int, contract_index: int):
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[0].paragraphs[0],
            self.borders("номер договора"),
            self.config["contracts_info"][contract_index][1]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[1].paragraphs[0],
            self.borders("период"),
            self.config["contracts_info"][contract_index][2]["contract_periods"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config["contracts_info"][contract_index][2]["accrual_debt"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[3].paragraphs[0],
            self.borders("неустойка"),
            self.config["contracts_info"][contract_index][2]["penalty"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[4].paragraphs[0],
            self.borders("неустойка+задолженность"),
            self.config["contracts_info"][contract_index][2]["debt_penalty"]
        )


    def third_table_fill_complex_row(self, table: Table, row_index: int, contract_index: int):
        self.third_table_merge_rows(table, row_index)

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[0].paragraphs[0],
            self.borders("номер договора"),
            self.config["contracts_info"][contract_index][1]
        )

        # Ячейка Период
        source_paragraph = table.row_cells(row_index)[1].paragraphs[0]
        new_paragraph = table.row_cells(row_index)[1].add_paragraph("")
        self.redactor.clone_paragraph(source_paragraph, new_paragraph)
        self.redactor.replace_text_in_paragraph(
            source_paragraph,
            self.borders("период"),
            self.config["contracts_info"][contract_index][2]["contract_periods"]
        )
        self.redactor.replace_text_in_paragraph(
            new_paragraph,
            self.borders("период"),
            "текущие начисления"
        )
        self.redactor.paragraph_text_set_bold(new_paragraph)

        source_paragraph = table.row_cells(row_index + 1)[1].paragraphs[0]
        new_paragraph = table.row_cells(row_index + 1)[1].add_paragraph("")
        self.redactor.clone_paragraph(source_paragraph, new_paragraph)
        self.redactor.replace_text_in_paragraph(
            source_paragraph,
            self.borders("период"),
            self.config["contracts_info"][contract_index][2]["contract_periods_correcting"]
        )
        self.redactor.replace_text_in_paragraph(
            new_paragraph,
            self.borders("период"),
            "доля от ГК"
        )
        self.redactor.paragraph_text_set_bold(new_paragraph)

        # Ячейка Задолженность
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config["contracts_info"][contract_index][2]["accrual_debt"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index + 1)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config["contracts_info"][contract_index][2]["correcting_debt"]
        )

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[3].paragraphs[0],
            self.borders("неустойка"),
            # "#неустойка#".upper()
            self.config["contracts_info"][contract_index][2]["penalty"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[4].paragraphs[0],
            self.borders("неустойка+задолженность"),
            # "#неустойка+задолженность#".upper()
            self.config["contracts_info"][contract_index][2]["debt_penalty"]
        )


    def third_table_fill_row_with_year_adjustment_only(self, table: Table, row_index: int, contract_index: int):
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[0].paragraphs[0],
            self.borders("номер договора"),
            self.config["contracts_info"][contract_index][1]
        )

        source_paragraph = table.row_cells(row_index)[1].paragraphs[0]
        new_paragraph = table.row_cells(row_index)[1].add_paragraph("")
        self.redactor.clone_paragraph(source_paragraph, new_paragraph)
        self.redactor.replace_text_in_paragraph(
            source_paragraph,
            self.borders("период"),
            self.config["contracts_info"][contract_index][2]["contract_periods_correcting"]
        )
        self.redactor.replace_text_in_paragraph(
            new_paragraph,
            self.borders("период"),
            "доля от ГК"
        )
        self.redactor.paragraph_text_set_bold(new_paragraph)

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config["contracts_info"][contract_index][2]["correcting_debt"]
        )

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[3].paragraphs[0],
            self.borders("неустойка"),
            # "#неустойка#".upper()
            self.config["contracts_info"][contract_index][2]["penalty"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[4].paragraphs[0],
            self.borders("неустойка+задолженность"),
            # "#неустойка+задолженность#".upper()
            self.config["contracts_info"][contract_index][2]["debt_penalty"]
        )


    def third_table_merge_rows(self, table: Table, row_index: int):
        row_1 = table.rows[row_index]
        row_2 = table.rows[row_index + 1]
        for i in [0, 3, 4]:
            self.redactor.merge_table_cells(row_1.cells[i], row_2.cells[i])


    def fill_first_list(self):
        """
        Создаем строку такого формата:
        Договорами (разделы 4, в том числе пункты 4.5 Договора № 02.104560-ТЭ от 16.05.2024, № 09.801771-ТЭ 
        от 23.08.2022, № 09.801895-ТЭ от 12.03.2024, а также раздел 5, в том числе п. 5.5 Договора № 
        09.801594-ТЭ от 01.04.2018)

        Сначала сортируем все договоры по разделам. В каждом разделе может быть один или несколько договоров.
        Если договор один, значит пишем <<раздел>>, если договоров несколько, пишем <<разделы>>. Если это 
        последний  раздел в списке разделов, и при этом не единственный, то добавляем <<а также>> перед 
        <<раздел ... >>. 
        Аналогично поступаем с пунктами. Внутри списка договоров с конкретным разделом разбиваем договоры на 
        разделы, учитываем единственное и множественное число, и фразу <<а также>>. 
        
        В итоге должна получиться структура:
        {
            'разделы': [
                {
                    'номер раздела': '1',
                    'кол-во договоров': 6,
                    'пункты': [
                        {'номер пункта': '1.1',
                         'договоры': ['№ 02.104560-ТЭ от 16.05.2024, № 09.801771-ТЭ от 23.08.2022', ...]},
                        {'номер пункта': '1.2',
                         'договоры': ['№ 02.104560-ТЭ от 16.05.2024, № 09.801771-ТЭ от 23.08.2022', ...]},
                        {'номер пункта': '1.3',
                         'договоры': ['№ 02.104560-ТЭ от 16.05.2024, № 09.801771-ТЭ от 23.08.2022', ...]},
                    ]
                },
                {
                    'номер раздела': '2',
                    'кол-во договоров': 6,
                    'пункты': [
                        {'номер пункта': '2.1',
                         'договоры': ['№ 02.104560-ТЭ от 16.05.2024, № 09.801771-ТЭ от 23.08.2022', ...]},
                        {'номер пункта': '2.2',
                         'договоры': ['№ 02.104560-ТЭ от 16.05.2024, № 09.801771-ТЭ от 23.08.2022', ...]},
                        {'номер пункта': '2.3',
                         'договоры': ['№ 02.104560-ТЭ от 16.05.2024, № 09.801771-ТЭ от 23.08.2022', ...]},
                    ]
                },
            ] 
        }
        """

        info = self.parse_info_for_first_list()
        text = self.create_text_for_first_list(info)

        template_text = "/*список разделов*/"
        start = self.redactor.find_paragraph_contains_text(template_text)

        self.redactor.replace_text_in_paragraph(
            self.redactor.get_paragraph(start),
            self.borders("список разделов"),
            text
        )


    def parse_info_for_first_list(self):
        info = {
            "chapters": []
        }
        parsed_chapters = []
        parsed_points = []
        rows_n = len(self.config["contracts_info"]) # Кол-во договоров
        for i in range(rows_n):
            point = self.config["contracts_info"][i][2]["contract_point"]
            chapter = point.split(".")[0]
            contract_number = self.config["contracts_info"][i][1]

            if chapter not in parsed_chapters:
                info["chapters"].append({
                    "number_of_chapter": chapter,
                    "number_of_contracts": 1,
                    "points": [{
                        "number_of_point": point,
                        "contracts": [contract_number]
                    }],
                })
                parsed_chapters.append(chapter)
                parsed_points.append(point)

            elif point not in parsed_points:
                for chapter_dict in info["chapters"]:
                    if chapter_dict["number_of_chapter"] == chapter:
                        chapter_dict["points"].append({
                            "number_of_point": point,
                            "contracts": [contract_number]
                        })
                        chapter_dict["number_of_contracts"] += 1
                        parsed_points.append(point)
                        break

            else:
                for chapter_dict in info["chapters"]:
                    if chapter_dict["number_of_chapter"] == chapter:
                        for point_dict in chapter_dict["points"]:
                            if point_dict["number_of_point"] == point:
                                point_dict["contracts"].append(contract_number)
                                chapter_dict["number_of_contracts"] += 1
                                break

        return info


    def create_text_for_first_list(self, info: dict):
        rows_n = len(self.config["contracts_info"]) # Кол-во договоров
        inserted_text = "("

        chapters_n = len(info["chapters"])
        for i, chapter_dict in enumerate(info["chapters"]):
            # Вставляем последний раздел
            if (i == chapters_n - 1) and (chapters_n > 1):
                inserted_text += "а также "

            inserted_text += "раздел"

            # обрабатываем случай с множественным числом
            if chapter_dict["number_of_contracts"] > 1:
                inserted_text += "ы"
            inserted_text += " "

            inserted_text += chapter_dict["number_of_chapter"]

            inserted_text += ", в том числе "

            points_n = len(chapter_dict["points"])
            for j, point_dict in enumerate(chapter_dict["points"]):
                if (j == points_n - 1) and (points_n > 1):
                    inserted_text += "а также "
                inserted_text += "пункт"

                if len(point_dict["contracts"]) > 1:
                    inserted_text += "ы"
                inserted_text += " "

                inserted_text += point_dict["number_of_point"]
                inserted_text += " Договора "

                contracts_n = len(point_dict["contracts"])
                for k, contract in enumerate(point_dict["contracts"]):
                    inserted_text += contract

                    if k < contracts_n - 1:
                        inserted_text += ", "

                if j < points_n - 1:
                    inserted_text += ", "

            if i < chapters_n - 1:
                inserted_text += ", "

        inserted_text += ")"

        return inserted_text


    def old_fill_first_list(self):
        rows_n = len(self.config["contracts_info"]) # Кол-во элементов в списке договоров
        template_text = f"раздел {self.borders('номер раздела')}, в том числе пункт {self.borders('пункт')} договора {self.borders('номер договора')}"
        start = self.redactor.find_paragraph_consists_of_text(template_text)   # Индекс первого абзаца списка

        for i in range(rows_n):
            if i < rows_n - 1:
                src_p = self.redactor.get_paragraph(start + i)
                new_p = self.redactor.insert_paragraph(start + i + 1)
                self.redactor.clone_paragraph(src_p, new_p)

            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i),
                self.borders("номер раздела"),
                # "#номер раздела#".upper()
                self.config["contracts_info"][i][2]["contract_point"].split(".")[0]
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i),
                self.borders("пункт"),
                self.config["contracts_info"][i][2]["contract_point"]
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i),
                self.borders("номер договора"),
                self.config["contracts_info"][i][1]
            )


    def fill_second_list(self):
        rows_n = len(self.config["lawsuit_info"]["claims"]) # Кол-во элементов в списке претензий
        template_text = f"list{self.borders('номер претензии')};"
        start = self.redactor.find_paragraph_consists_of_text(template_text)   # Индекс первого абзаца списка

        for i in range(rows_n):
            if i < rows_n - 1:
                src_p = self.redactor.get_paragraph(start + i)
                new_p = self.redactor.insert_paragraph(start + i + 1)
                self.redactor.clone_paragraph(src_p, new_p)

            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i),
                template_text[:len(template_text) - 1],
                self.config["lawsuit_info"]["claims"][i]
            )


    def fill_third_list(self):
        rows_n = len(self.config["contracts_info"]) # Кол-во элементов в списке договоров
        template_text = "по Договору 05.403297-ТЭ от 01.08.2017:"
        start = self.redactor.find_paragraph_consists_of_text(template_text)   # Индекс первого абзаца списка

        # Список индексов тех параграфов, которые нужно удалить после цикла. Почему их нужно удалять?
        # Эти параграфы - это корректировка обязательств, которая есть не у всех договоров. Поэтому 
        # там где ее нет, параграфы удаляются
        delete_par_indices = []

        # Эта переменная нужна для подсчета числа заполненных строк в списке. Нельзя просто сделать
        # rows_n * stride, поскольку у некоторых договоров отсутствует годовая корректировка, и 
        # соответствующие параграфы удаляются из списка, т.е. строк становится меньше
        filled_rows = 0

        stride = 5
        for i in range(rows_n):
            if i < rows_n - 1:
                for j in range(stride):
                    src_p = self.redactor.get_paragraph(start + i * stride + j)
                    new_p = self.redactor.insert_paragraph(start + (i + 1) * stride + j)
                    self.redactor.clone_paragraph(src_p, new_p)

            # contract_number = list(self.config["table_info"].keys())[i]
            contract_number = self.config["contracts_info"][i][1]

            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride),
                "05.403297-ТЭ от 01.08.2017",
                # self.config["contracts_info"][i][1]
                contract_number
            )
            
            if self.config["table_info"][contract_number]["accrual_debt"] is not None:
                if self.config["table_info"][contract_number]["accrual_debt"] != "0,00":
                    self.redactor.replace_text_in_paragraph(
                        self.redactor.get_paragraph(start + i * stride + 1),
                        self.borders("задолженность"),
                        # self.config["contracts_info"][i][2]["debt"]
                        self.config["table_info"][contract_number]["accrual_debt"]
                    )
                    self.redactor.replace_text_in_paragraph(
                        self.redactor.get_paragraph(start + i * stride + 1),
                        self.borders("период"),
                        # self.config["contracts_info"][i][2]["contract_periods"]
                        self.config["table_info"][contract_number]["contract_periods"]  
                    )
                else:
                    if start + i * stride + 1 not in delete_par_indices:
                        delete_par_indices.append(start + i * stride + 1)
            else:
                if start + i * stride + 1 not in delete_par_indices:
                    delete_par_indices.append(start + i * stride + 1)

            if self.config["table_info"][contract_number]["contract_periods_correcting"] is not None:
                if self.config["table_info"][contract_number]["correcting_debt"] != "0,00":
                    self.redactor.replace_text_in_paragraph(
                        self.redactor.get_paragraph(start + i * stride + 2),
                        self.borders("доля годовой корректировки"),
                        self.config["table_info"][contract_number]["correcting_debt"]
                    )
                    self.redactor.replace_text_in_paragraph(
                        self.redactor.get_paragraph(start + i * stride + 2),
                        self.borders("период корректировки"),
                        self.config["table_info"][contract_number]["contract_periods_correcting"]
                    )
                    self.redactor.replace_text_in_paragraph(
                        self.redactor.get_paragraph(start + i * stride + 2),
                        self.borders("год корректировки"),
                        self.config["table_info"][contract_number]["correcting_year"]
                    )
                else:
                    if start + i * stride + 2 not in delete_par_indices:
                        delete_par_indices.append(start + i * stride + 2)
            else:
                if start + i * stride + 2 not in delete_par_indices:
                    delete_par_indices.append(start + i * stride + 2)

            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride + 3),
                self.borders("неустойка"),
                # self.config["contracts_info"][i][2]["debt_penalty"]
                self.config["table_info"][contract_number]["penalty"]
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride + 3),
                self.borders("период неустойки 1"),
                # self.config["contracts_info"][i][2]["debt_penalty"]
                self.config["table_info"][contract_number]["penalty_period"].split(" по ")[0]
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride + 3),
                self.borders("период неустойки 2"),
                # self.config["contracts_info"][i][2]["debt_penalty"]
                self.config["table_info"][contract_number]["penalty_period"].split(" по ")[1]
            )

            date_string = self.config["table_info"][contract_number]["penalty_period"].split(" по ")[1]
            format_string = "%d.%m.%Y"
            parsed_date = datetime.strptime(date_string, format_string)
            parsed_date_plus_1 = parsed_date + timedelta(days=1)
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride + 4),
                self.borders("начальная дата неустойки"),
                # self.config["contracts_info"][i][2]["debt_penalty"]
                parsed_date_plus_1.strftime(format_string)
            )

        delete_paragraps = 0
        for index in delete_par_indices:
            self.redactor.delete_paragraph(index - delete_paragraps)
            delete_paragraps += 1


    def fill_other_parts(self):
        # Добавление нужного параграфа я решил реализовать отдельно
        # Проверка > 0 нужна, потому что этот код учитывает только ТЭ и ГВС договоры,
        # на других (например ФОТЭ) будет крашиться
        new_paragraph = None
        significant_paragraph_template = "/*абзац важный*/"
        significant_paragraph = self.redactor.get_paragraph(
            self.redactor.find_paragraph_contains_text(significant_paragraph_template)
        )
        paragraphs_need_to_be_inserted = [significant_paragraph]

        if len(self.config["contract_types_templates"]["types_of_significant_paragraph"]) > 0:

            if len(self.config["contract_types_templates"]["types_of_significant_paragraph"]) == 2:
                new_paragraph = self.redactor.insert_paragraph_after_paragraph(significant_paragraph)
                self.redactor.clone_paragraph(significant_paragraph, new_paragraph)
                paragraphs_need_to_be_inserted.append(new_paragraph)

            for i, para in enumerate(paragraphs_need_to_be_inserted):
                for run in para.runs:
                    if significant_paragraph_template in run.text:
                        inserting_text = self.get_paragraph_type_of(self.config["contract_types_templates"]["types_of_significant_paragraph"][i])
                        run.text = run.text.replace(significant_paragraph_template, inserting_text)
        else:
            # self.redactor.delete_paragraph()
            pass


        # # Эту штуку просто не получается запихать в calculator_adapter
        # # потому что там еще нет service_type
        # converted_data["lawsuit_info"]["service_type"]


        to_replace = {
            "/*цена иска*/": self.config["lawsuit_info"]["cost"],
            "/*госпошлина*/": self.config["lawsuit_info"]["tax"],
            "/*истец полное имя*/": self.config["plaintiff_info"]["full_name"],
            "/*истец сокращенное имя*/": self.config["plaintiff_info"]["short_name"],
            "/*ответчик полное имя*/": self.config["defendant_info"]["full_name"].upper(),
            "/*ответчик сокращенное имя*/": self.config["defendant_info"]["short_name"],
            "/*тип договора*/": self.config["contract_types_templates"]["contract_type"],
            # "/*тип договора2*/": self.config["contract_types_templates"]["contract_type2"],
            "/*сумма долга*/": self.config["table_info"]["all_debt"],
            "/*истец огрн*/": self.config["plaintiff_info"]["ogrn"],
            "/*истец инн*/": self.config["plaintiff_info"]["inn"],
            "/*ответчик огрн*/": self.config["defendant_info"]["ogrn"],
            "/*ответчик инн*/": self.config["defendant_info"]["inn"],
            "/*поставляемые ресурсы*/": self.config["contract_types_templates"]["supplied_resources"],
            "/*поставляемые ресурсы2*/": self.config["contract_types_templates"]["supplied_resources2"],
            "/*поставляемые ресурсы3*/": self.config["contract_types_templates"]["supplied_resources3"],
            "/*поставляемые ресурсы4*/": self.config["contract_types_templates"]["supplied_resources4"],
            "/*поставляемые ресурсы5*/": self.config["contract_types_templates"]["supplied_resources5"],
            "/*мн.ч.*/": self.config["contract_types_templates"]["plural_template_1"],
            "/*мн.ч.2*/": self.config["contract_types_templates"]["plural_template_2"],
            "/*мн.ч.3*/": self.config["contract_types_templates"]["plural_template_3"],
            "/*мн.ч.4*/": self.config["contract_types_templates"]["plural_template_4"],
            "/*мн.ч.5*/": self.config["contract_types_templates"]["plural_template_5"],
            "/*мн.ч.6*/": self.config["contract_types_templates"]["plural_template_6"],
            "/*мн.ч.7*/": self.config["contract_types_templates"]["plural_template_7"],
            "/*нужная статья*/": self.config["contract_types_templates"]["service_article"],

            "/*направлен окончание*/": self.config["contract_types_templates"]["directed_ending"],
            "/*претензия окончание и.п*/": self.config["contract_types_templates"]["claim_ending_1"],
            "/*которая окончание*/": self.config["contract_types_templates"]["which_ending"],
            "/*оставлен окончание*/": self.config["contract_types_templates"]["abandoned_ending"],
            "/*претензии окончание р.п*/": self.config["contract_types_templates"]["claim_ending_2"],
            "/*её число*/": self.config["contract_types_templates"]["her_ending"],
            "/*копия окончание*/": self.config["contract_types_templates"]["copy_ending"],
        }
        for paragraph in self.doc.paragraphs:
            # runs = []
            for run in paragraph.runs:
                # runs.append(run.text)
                for k, v in to_replace.items():
                    if k in run.text:
                        run.text = run.text.replace(k, v)
            # print(runs)


    def get_paragraph_text_6(self):
        pass


    def get_paragraph_type_of(self, str_type: str):
        if "ГВС" in str_type:
            if "ТСЖ" in str_type:
                return self.__paragraph_type_2()
            elif "Прочие".lower() in str_type.lower():
                return self.__paragraph_type_4()
            elif "УК" in str_type:
                return self.__paragraph_type_6()
            else:
                print(f"Тип компании str_type={str_type} не соответствует ни одному формату, не ясно как обрабатывать.")
                raise RuntimeError(f"Invalid str_type: {str_type}")

        elif "ТЭ" in str_type:
            if "ТСЖ" in str_type:
                return self.__paragraph_type_1()
            elif "Прочие".lower() in str_type.lower():
                return self.__paragraph_type_3()
            elif "УК" in str_type:
                return self.__paragraph_type_5()
            else:
                print(f"Тип компании str_type={str_type} не соответствует ни одному формату, не ясно как обрабатывать.")
                raise RuntimeError(f"Invalid str_type: {str_type}")

        else:
            print(f"Тип договора str_type={str_type} не соответствует ни одному формату, не ясно как обрабатывать.")
            raise RuntimeError(f"Invalid str_type: {str_type}")


    def __paragraph_type_1(self):
        "ТСЖ + ТЭ"
        return "В соответствии с частью 9.2 статьи 15 Федерального закона от 27.07.2010 № 190-ФЗ «О теплоснабжении» товарищества собственников жилья, жилищные, жилищно-строительные и иные специализированные потребительские кооперативы, созданные в целях удовлетворения потребностей граждан в жилье, приобретающие тепловую энергию (мощность) и (или) теплоноситель для целей предоставления коммунальных услуг, в случае несвоевременной и (или) неполной оплаты тепловой энергии (мощности) и (или) теплоносителя уплачивают единой теплоснабжающей организации (теплоснабжающей организации) пени в размере одной трехсотой ставки рефинансирования Центрального банка Российской Федерации, действующей на день фактической оплаты, от не выплаченной в срок суммы за каждый день просрочки начиная с тридцать первого дня, следующего за днем наступления установленного срока оплаты, по день фактической оплаты, произведенной в течение девяноста календарных дней со дня наступления установленного срока оплаты, либо до истечения девяноста календарных дней после дня наступления установленного срока оплаты, если в девяностодневный срок оплата не произведена. Начиная с девяносто первого дня, следующего за днем наступления установленного срока оплаты, по день фактической оплаты пени уплачиваются в размере одной стотридцатой ставки рефинансирования Центрального банка Российской Федерации, действующей на день фактической оплаты, от не выплаченной в срок суммы за каждый день просрочки."
    
    
    def __paragraph_type_2(self):
        "ТСЖ + ГВС"
        return "Согласно части 6.3 статьи 13 Федерального закона от 07.12.2011 № 416-ФЗ «О водоснабжении и водоотведении» товарищества собственников жилья, жилищные, жилищно-строительные и иные специализированные потребительские кооперативы, созданные в целях удовлетворения потребностей граждан в жилье, приобретающие горячую, питьевую и (или) техническую воду для целей предоставления коммунальных услуг, в случае несвоевременной и (или) неполной оплаты воды уплачивают организации, осуществляющей горячее водоснабжение, холодное водоснабжение, пени в размере одной трехсотой ставки рефинансирования Центрального банка Российской Федерации, действующей на день фактической оплаты, от не выплаченной в срок суммы за каждый день просрочки начиная с тридцать первого дня, следующего за днем наступления установленного срока оплаты, по день фактической оплаты, произведенной в течение девяноста календарных дней со дня наступления установленного срока оплаты, либо до истечения девяноста календарных дней после дня наступления установленного срока оплаты, если в девяностодневный срок оплата не произведена. Начиная с девяносто первого дня, следующего за днем наступления установленного срока оплаты, по день фактической оплаты пени уплачиваются в размере одной стотридцатой ставки рефинансирования Центрального банка Российской Федерации, действующей на день фактической оплаты, от не выплаченной в срок суммы за каждый день просрочки."
    
    
    def __paragraph_type_3(self):
        "Прочие + ТЭ"
        return "В соответствии с частью 9.1 статьи 15 Федерального закона от 27.07.2010 № 190-ФЗ «О теплоснабжении» потребитель тепловой энергии, несвоевременно и (или) не полностью оплативший тепловую энергию (мощность) и (или) теплоноситель по договору теплоснабжения, обязан уплатить единой теплоснабжающей организации (теплоснабжающей организации) пени в размере одной стотридцатой ставки рефинансирования Центрального банка Российской Федерации, действующей на день фактической оплаты, от не выплаченной в срок суммы за каждый день просрочки начиная со следующего дня после дня наступления установленного срока оплаты по день фактической оплаты."
    
    
    def __paragraph_type_4(self):
        "Прочие + ГВС"
        return "Согласно части 6.2 статьи 13 Федерального закона от 07.12.2011 № 416-ФЗ «О водоснабжении и водоотведении» абонент, несвоевременно и (или) не полностью оплативший горячую, питьевую и (или) техническую воду, обязан уплатить организации, осуществляющей горячее водоснабжение, холодное водоснабжение, пени в размере одной стотридцатой ставки рефинансирования Центрального банка Российской Федерации, действующей на день фактической оплаты, от не выплаченной в срок суммы за каждый день просрочки начиная со следующего дня после дня наступления установленного срока оплаты по день фактической оплаты."
    
    
    def __paragraph_type_5(self):
        "УК + ТЭ"
        return "В соответствии с частью 9.3 статьи 15 Федерального закона от 27.07.2010 № 190-ФЗ «О теплоснабжении» управляющие организации, приобретающие тепловую энергию (мощность) и (или) теплоноситель для целей предоставления коммунальных услуг, организации, осуществляющие горячее водоснабжение, холодное водоснабжение и (или) водоотведение по договорам горячего водоснабжения и договорам поставки горячей воды, а также теплоснабжающие организации, приобретающие тепловую энергию (мощность) и (или) теплоноситель по договору поставки тепловой энергии (мощности) и (или) теплоносителя, в случае несвоевременной и (или) неполной оплаты тепловой энергии (мощности) и (или) теплоносителя уплачивают единой теплоснабжающей организации (теплоснабжающей организации) пени в размере одной трехсотой ставки рефинансирования Центрального банка Российской Федерации, действующей на день фактической оплаты, от не выплаченной в срок суммы за каждый день просрочки начиная со дня, следующего за днем наступления установленного срока оплаты, по день фактической оплаты, произведенной в течение шестидесяти календарных дней со дня наступления установленного срока оплаты, либо до истечения шестидесяти календарных дней после дня наступления установленного срока оплаты, если в шестидесятидневный срок оплата не произведена. Начиная с шестьдесят первого дня, следующего за днем наступления установленного срока оплаты, по день фактической оплаты, произведенной в течение девяноста календарных дней со дня наступления установленного срока оплаты, либо до истечения девяноста календарных дней после дня наступления установленного срока оплаты, если в девяностодневный срок оплата не произведена, пени уплачиваются в размере одной стосемидесятой ставки рефинансирования Центрального банка Российской Федерации, действующей на день фактической оплаты, от не выплаченной в срок суммы за каждый день просрочки. Начиная с девяносто первого дня, следующего за днем наступления установленного срока оплаты, по день фактической оплаты пени уплачиваются в размере одной стотридцатой ставки рефинансирования Центрального банка Российской Федерации, действующей на день фактической оплаты, от не выплаченной в срок суммы за каждый день просрочки."
    
    
    def __paragraph_type_6(self):
        "УК + ГВС"
        return "Согласно части 6.4 статьи 13 Федерального закона от 07.12.2011 № 416-ФЗ «О водоснабжении и водоотведении» управляющие организации, приобретающие горячую, питьевую и (или) техническую воду для целей предоставления коммунальных услуг, теплоснабжающие организации (единые теплоснабжающие организации), а также организации, осуществляющие горячее водоснабжение, холодное водоснабжение, приобретающие горячую, питьевую и (или) техническую воду по договорам горячего водоснабжения, договорам холодного водоснабжения или единым договорам холодного водоснабжения и водоотведения, в случае несвоевременной и (или) неполной оплаты горячей, питьевой и (или) технической воды уплачивают организации, осуществляющей горячее водоснабжение, холодное водоснабжение, пени в размере одной трехсотой ставки рефинансирования Центрального банка Российской Федерации, действующей на день фактической оплаты, от не выплаченной в срок суммы за каждый день просрочки начиная со дня, следующего за днем наступления установленного срока оплаты, по день фактической оплаты, произведенной в течение шестидесяти календарных дней со дня наступления установленного срока оплаты, либо до истечения шестидесяти календарных дней после дня наступления установленного срока оплаты, если в шестидесятидневный срок оплата не произведена. Начиная с шестьдесят первого дня, следующего за днем наступления установленного срока оплаты, по день фактической оплаты, произведенной в течение девяноста календарных дней со дня наступления установленного срока оплаты, либо до истечения девяноста календарных дней после дня наступления установленного срока оплаты, если в девяностодневный срок оплата не произведена, пени уплачиваются в размере одной стосемидесятой ставки рефинансирования Центрального банка Российской Федерации, действующей на день фактической оплаты, от не выплаченной в срок суммы за каждый день просрочки. Начиная с девяносто первого дня, следующего за днем наступления установленного срока оплаты, по день фактической оплаты пени уплачиваются в размере одной стотридцатой ставки рефинансирования Центрального банка Российской Федерации, действующей на день фактической оплаты, от не выплаченной в срок суммы за каждый день просрочки."
    
    
    
    
    def fix_quotes(self):
        self.config['plaintiff_info']['full_name'] = self.normalize_quotes(self.config['plaintiff_info']['full_name'])
        self.config['plaintiff_info']['short_name'] = self.normalize_quotes(self.config['plaintiff_info']['short_name'])
        self.config['defendant_info']['full_name'] = self.normalize_quotes(self.config['defendant_info']['full_name'])
        self.config['defendant_info']['short_name'] = self.normalize_quotes(self.config['defendant_info']['short_name'])

    def normalize_quotes(self, text: str) -> str:
        # Список всех возможных кавычек, которые нужно заменить
        quote_chars = {'\'', '"', '«', '»', '„', '“', '”'}
        
        # Преобразуем строку в список для удобства замены по индексу
        chars = list(text)
        n = len(chars)
        
        i = 0
        while i < n:
            if chars[i] in quote_chars:
                # Проверяем по правилам в порядке приоритета
                
                # 1) В начале строки
                if i == 0:
                    chars[i] = '«'
                # 2) В конце строки
                elif i == n - 1:
                    chars[i] = '»'
                # 3) Перед кавычкой пробел
                elif chars[i - 1] == ' ':
                    chars[i] = '«'
                # 4) После кавычки пробел
                elif chars[i + 1] == ' ':
                    chars[i] = '»'
                # 5) Во всех остальных случаях
                else:
                    chars[i] = '«'
            i += 1
        
        return ''.join(chars)


    def quote_type(self, quote: str, quotes: list) -> tuple:
        """
        Воозвращает кортеж из двух значений: (индекс кортежа с кавчиками, открытая/закрытая)
        1 - открывающая
        -1 - закрывающая
        """
        for i, item in enumerate(quotes):
            if quote == item[0]:
                return (i, 1)   # Открывающая
            elif quote == item[1]:
                return (i, -1)    # Закрывающая
        return None


    def prepare_data(self):
        """
        Это вспомогательная функция. Она нужна чтобы определить, какой иск генерировать. 
        Сейчас есть три вида: ТЭ, ГВС и ТЭ + ГВС. В зависимости от этого, используются 
        разные шаблоны абзацев.
        """
        templates = dict()
        contracts = self.config["contracts_info"]
        company_type = self.config["company_type"]
        service_type = self.config["lawsuit_info"]["service_type"]

        contruct_types = []
        if "ГВС" in service_type:
            contruct_types.append("ГВС")
        if "ТЭ" in service_type:
            contruct_types.append("ТЭ")

        if len(contracts) == 1:
            templates["supplied_resources4"] = "ТЭ"
            templates["plural_template_1"] = "Договором"
            templates["plural_template_2"] = "Договору"
            templates["plural_template_3"] = "названном Договоре"
            templates["plural_template_4"] = "Договора"
            templates["plural_template_5"] = "указанного Договора"
            templates["plural_template_6"] = "названному Договору"
            templates["plural_template_7"] = "был заключен следующий договор (далее именуемый – Договор), предметом которого"

        elif len(contracts) > 1:
            templates["plural_template_1"] = "Договорами"
            templates["plural_template_2"] = "Договорам"
            templates["plural_template_3"] = "названных Договорах"
            templates["plural_template_4"] = "Договоров"
            templates["plural_template_5"] = "указанных Договоров"
            templates["plural_template_6"] = "названным Договорам"
            templates["plural_template_7"] = "были заключены следующие договоры (далее именуемые – Договоры), предметом которых"


        templates["types_of_significant_paragraph"] = []
        # Выбираем нужные шаблоны
        if ("ГВС" in contruct_types) and ("ТЭ" in contruct_types):
            templates["supplied_resources"] = "тепловой энергии и/или теплоносителя (далее – ТЭ), горячей воды через присоединенные сети горячего водоснабжения (далее – ГВС)"
            templates["contract_type"] = "тепловую энергию/теплоноситель (ТЭ) и горячую воду (ГВС)"
            # templates["contract_type2"] = "тепловую энергию/теплоноситель (ТЭ) и горячую воду (ГВС)"
            templates["supplied_resources2"] = "тепловой энергии/теплоносителя, горячей воды"
            templates["supplied_resources3"] = "тепловую энергию/теплоноситель, горячую воду"
            templates["supplied_resources4"] = "ТЭ и ГВС"
            templates["supplied_resources5"] = "тепловую энергию и поставку горячей воды"
            templates["types_of_significant_paragraph"].append(company_type + "ТЭ")
            templates["types_of_significant_paragraph"].append(company_type + "ГВС")
            templates["service_article"] = "ст. 15 Федерального закона от 27.07.2010 № 190-ФЗ «О теплоснабжении», ст. 13 Федерального закона от 07.12.2011 № 416-ФЗ «О водоснабжении и водоотведении»"


        elif "ТЭ" in contruct_types:
            templates["supplied_resources"] = "тепловой энергии и/или теплоносителя (далее – ТЭ)"
            templates["contract_type"] = "тепловую энергию/теплоноситель (ТЭ)"
            # templates["contract_type2"] = "тепловую энергию/теплоноситель (ТЭ)"
            templates["supplied_resources2"] = "тепловой энергии/теплоносителя"
            templates["supplied_resources3"] = "тепловую энергию/теплоноситель"
            templates["supplied_resources4"] = "ТЭ"
            templates["supplied_resources5"] = "тепловую энергию"
            templates["types_of_significant_paragraph"].append(company_type + "ТЭ")
            templates["service_article"] = "ст. 15 Федерального закона от 27.07.2010 № 190-ФЗ «О теплоснабжении»"

        elif "ГВС" in contruct_types:
            templates["supplied_resources"] = "горячей воды через присоединенные сети горячего водоснабжения (далее – ГВС)"
            templates["contract_type"] = "горячую воду (ГВС)"
            # templates["contract_type2"] = "горячую воду (ГВС)"
            templates["supplied_resources2"] = "горячей воды"
            templates["supplied_resources3"] = "горячую воду"
            templates["supplied_resources4"] = "ГВС"
            templates["supplied_resources5"] = "поставку горячей воды"
            templates["types_of_significant_paragraph"].append(company_type + "ГВС")
            templates["service_article"] = "ст. 13 Федерального закона от 07.12.2011 № 416-ФЗ «О водоснабжении и водоотведении»"


        number_of_claims = len(self.config["lawsuit_info"]["claims"]) 
        if number_of_claims == 1:
            templates["directed_ending"] = "а"
            templates["claim_ending_1"] = "я"
            templates["which_ending"] = "ая"
            templates["abandoned_ending"] = "а"
            templates["claim_ending_2"] = "и"
            templates["her_ending"] = "её"
            templates["copy_ending"] = "я"
        elif number_of_claims > 1:
            templates["directed_ending"] = "ы"
            templates["claim_ending_1"] = "и"
            templates["which_ending"] = "ые"
            templates["abandoned_ending"] = "ы"
            templates["claim_ending_2"] = "й"
            templates["her_ending"] = "их"
            templates["copy_ending"] = "и"
        else:
            print(f"В списке претензий в данных нет ни одной претензии")
            raise RuntimeError(f"No claims in config data: {self.config['lawsuit_info']}")

        self.config["contract_types_templates"] = templates

        # if self.config["defendant_info"]["full_name"].lower().startswith("ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ".lower()):
        #     self.config["defendant_info"]["full_name"] = "Общество с ограниченной ответственностью"
        # self.config["defendant_info"]["full_name"] = self.config["defendant_info"]["full_name"].replace()

        _pattern = r"ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ"
        _replacement = "Общество с ограниченной ответственностью"
        text = self.config["defendant_info"]["full_name"]
        self.config["defendant_info"]["full_name"] = re.sub(_pattern, _replacement, text, flags=re.IGNORECASE)



        # # УБРАТЬ ЭТУ ЗАГЛУШКУ
        # rows_n = len(self.config["contracts_info"])
        # for i in range(rows_n):
        #     contract_number = self.config["contracts_info"][i][1]
        #     self.config["table_info"][contract_number]["correcting_year"] = "2024"
