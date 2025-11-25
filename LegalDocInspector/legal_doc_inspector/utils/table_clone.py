from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy

def get_cell_tcW(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is not None:
        return int(tcW.get(qn('w:w')))
    return None


def get_cell_shading(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        return None
    return {
        "fill": shd.get(qn("w:fill")),
        "color": shd.get(qn("w:color")),
        "val": shd.get(qn("w:val")),
    }


def get_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        return None

    result = {}
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = borders.find(qn(f"w:{side}"))
        if tag is not None:
            result[side] = {
                "val": tag.get(qn("w:val")),
                "sz": tag.get(qn("w:sz")),
                "color": tag.get(qn("w:color")),
            }
    return result


def parse_table_full_structure(table):
    """
    Извлекает:
    - структуру строк/ячееек
    - текст
    - форматирование параграфов
    - объединения
    - ширины ячеек
    - высоту строк
    - заливки
    - borders
    - tblGrid
    """
    structure = {
        "rows": [],
        "tblGrid": [],
        "style": table.style.name if table.style else None
    }

    # --- tblGrid ---
    if table._element.tblGrid is not None:
        for col in table._element.tblGrid.iter(qn("w:gridCol")):
            structure["tblGrid"].append(int(col.get(qn("w:w"))))

    # --- parsing rows ---
    for r_idx, row in enumerate(table.rows):
        row_info = {
            "height": None,
            "cells": []
        }

        # высота строки
        trPr = row._tr.get_or_add_trPr()
        trHeight = trPr.find(qn("w:trHeight"))
        if trHeight is not None:
            row_info["height"] = int(trHeight.get(qn("w:val")))

        # parsing cells
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            grid_span = tcPr.find(qn('w:gridSpan'))
            grid_span = int(grid_span.get(qn('w:val'))) if grid_span is not None else 1

            v_merge = tcPr.find(qn('w:vMerge'))
            if v_merge is not None:
                v_merge = v_merge.get(qn('w:val')) or "continue"

            cell_info = {
                "source_cell": cell,
                "text": cell.text,
                "gridSpan": grid_span,
                "vMerge": v_merge,
                "width": get_cell_tcW(cell),
                "shading": get_cell_shading(cell),
                "borders": get_cell_borders(cell),
            }

            row_info["cells"].append(cell_info)

        structure["rows"].append(row_info)

    return structure



def apply_cell_formatting(cell, info):
    """Применяет ширину, shading, borders."""
    tcPr = cell._tc.get_or_add_tcPr()

    # ширина
    if info["width"]:
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(info["width"]))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    # заливка
    if info["shading"]:
        shd = OxmlElement("w:shd")
        for k, v in info["shading"].items():
            shd.set(qn(f"w:{k}"), v)
        tcPr.append(shd)

    # borders
    if info["borders"]:
        borders = OxmlElement("w:tcBorders")
        for side, data in info["borders"].items():
            tag = OxmlElement(f"w:{side}")
            for k, v in data.items():
                if v:
                    tag.set(qn(f"w:{k}"), v)
            borders.append(tag)
        tcPr.append(borders)


def clone_table_full(document, structure):
    """Создаёт таблицу с полным копированием стилей."""
    rows_count = len(structure["rows"])

    # столбцы определяем по tblGrid
    if structure["tblGrid"]:
        cols_count = len(structure["tblGrid"])
    else:
        # fallback — максимальное количество логических ячеек
        cols_count = max(sum(c["gridSpan"] for c in row["cells"]) for row in structure["rows"])

    new_table = document.add_table(rows=rows_count, cols=cols_count)

    # применяем стиль таблицы
    if structure["style"]:
        new_table.style = structure["style"]

    # создаём tblGrid
    if structure["tblGrid"]:
        tbl = new_table._element
        if tbl.tblGrid is not None:
            tbl.remove(tbl.tblGrid)

        new_grid = OxmlElement("w:tblGrid")
        for w in structure["tblGrid"]:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(w))
            new_grid.append(col)
        tbl.insert(0, new_grid)

    # копирование содержимого
    for r_idx, row_info in enumerate(structure["rows"]):
        row = new_table.rows[r_idx]

        # высота строки
        if row_info["height"]:
            trPr = row._tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), str(row_info["height"]))
            trHeight.set(qn("w:hRule"), "exact")
            trPr.append(trHeight)

        col_idx = 0
        for cell_info in row_info["cells"]:
            # Проверка, не выйдем ли за пределы row.cells
            if col_idx >= len(row.cells):
                break

            cell = row.cells[col_idx]

            # текст
            copy_cell_text_with_styles(cell_info['source_cell'], cell)

            # gridSpan
            if cell_info["gridSpan"] > 1:
                gridSpan = OxmlElement('w:gridSpan')
                gridSpan.set(qn('w:val'), str(cell_info["gridSpan"]))
                cell._tc.get_or_add_tcPr().append(gridSpan)

            # vMerge
            if cell_info.get("vMerge"):
                vMerge = OxmlElement("w:vMerge")
                if cell_info["vMerge"] != "continue":
                    vMerge.set(qn("w:val"), cell_info["vMerge"])
                cell._tc.get_or_add_tcPr().append(vMerge)

            # форматирование
            apply_cell_formatting(cell, cell_info)

            col_idx += cell_info["gridSpan"]  # корректно учитываем объединённые ячейки

    return new_table



def copy_paragraph_formatting(src_p, dst_p):
    """Копирует форматирование параграфа (alignment, spacing, indentation)."""

    # выравнивание
    dst_p.alignment = src_p.alignment

    # spacing + indentation (копируем XML напрямую)
    pPr = src_p._p.get_or_add_pPr()
    new_pPr = dst_p._p.get_or_add_pPr()

    for child in pPr:
        tag = child.tag
        # не копируем rPr (форматирование текста), только свойства параграфа
        if tag.endswith("rPr"):
            continue
        new_pPr.append(deepcopy(child))


def copy_run_formatting(src_run, dst_run):
    """Копирует все свойства run'а из src_run в dst_run."""

    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.strike = src_run.font.strike
    dst_run.font.superscript = src_run.font.superscript
    dst_run.font.subscript = src_run.font.subscript

    # цвет
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb

    # имя шрифта
    if src_run.font.name:
        dst_run.font.name = src_run.font.name

    # размер
    if src_run.font.size:
        dst_run.font.size = src_run.font.size


def copy_cell_text_with_styles(src_cell, dst_cell):
    """
    Копирует текст и форматирование из src_cell → dst_cell
    включая параграфы, run'ы, их форматирование.
    """
    dst_cell.text = ""  # очищаем по умолчанию

    dst_cell_paragraphs = dst_cell.paragraphs
    # удаляем пустые созданные параграфы
    for p in dst_cell_paragraphs:
        p._element.getparent().remove(p._element)

    for src_paragraph in src_cell.paragraphs:
        dst_paragraph = dst_cell.add_paragraph()
        copy_paragraph_formatting(src_paragraph, dst_paragraph)

        for src_run in src_paragraph.runs:
            dst_run = dst_paragraph.add_run(src_run.text)
            copy_run_formatting(src_run, dst_run)

def clone_table_from_table(document, table):
    """
    Создаёт полную копию таблицы (со структурой, стилями, текстом, форматами)
    и возвращает объект Table, НЕ вставляя его автоматически в документ.
    """
    struct = parse_table_full_structure(table)

    # Создаём временную таблицу в документе
    new_table = clone_table_full(document, struct)

    # Удаляем её из документа, чтобы не вставлялась автоматически
    tbl_element = new_table._element
    body = tbl_element.getparent()
    body.remove(tbl_element)

    # Теперь new_table можно вставлять куда угодно вручную
    return new_table
