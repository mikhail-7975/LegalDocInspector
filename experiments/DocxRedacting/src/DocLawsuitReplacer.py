import json

from docx import Document
from docx.table import Table

from DocxRedactor import DocxRedactor


class DocLawsuitReplacer:
    """
    Класс для редактирования документа с табличкой
    
    """
    def __init__(self) -> None:
        self.redactor: DocxRedactor = DocxRedactor()
        self.config: dict = None
        self.doc: Document = None


    def make_instance(self, config_filename: str, template_filename: str, output_filename: str = None):
        """
        config_filename: str - имя файла с конфигами для шаблона
        template_filename: str - имя файла-шаблона
        output_filename: str - имя файла, созданного по шаблону
        """
        self.config = self.parse_config(config_filename)

        self.redactor.clone_file(template_filename, output_filename)
        self.doc = self.redactor.open(output_filename)

        self.fill_file()

        self.redactor.save()
        self.redactor.close()


    def parse_config(self, filename: str) -> dict:
        with open(filename, 'r', encoding="utf-8") as file:
            data = json.load(file)
        return data


    def fill_file(self):
        self.clone_table(len(self.config["contracts"]) - 1)
        
        self.redactor.print_table(self.redactor.get_table(0))

        for i, contract in enumerate(self.config["contracts"]):
            self.fill_contract(contract, i)


    def fill_contract(self, contract, contract_index):
        table = self.redactor.get_table(contract_index)
        
        self.fill_common_contract_info(contract, table)

        self.clone_block(table, 4, len(contract["periods"]) - 1)

        payments = 0
        for i, period in enumerate(contract["periods"]):
            self.fill_period(period, i, table, payments)
            payments += len(period["payments_1"])
            payments += len(period["payments_2"])


    def fill_common_contract_info(self, contract, table: Table):
        to_replace = (
            (self.borders("номер договора"),            contract["contract_number"]),
            (self.borders("дата начала просрочки"),     contract["start_date_of_delay"]),
            (self.borders("дата конца просрочки"),      contract["end_date_of_delay"]),
            (self.borders("сумма долга"),               contract["sum_debt"]),
            (self.borders("сумма пеней"),               contract["sum_peny"]),
        )

        cell = table.row_cells(0)[2]
        flag = self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])
        cell = table.row_cells(1)[2]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[1][0], to_replace[1][1])
        cell = table.row_cells(1)[10]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[2][0], to_replace[2][1])
        cell = table.row_cells(9)[0]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[3][0], to_replace[3][1])
        cell = table.row_cells(10)[0]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[4][0], to_replace[4][1])


    def fill_period(self, period, period_index, table: Table, payments: int):
        self.fill_common_period_info(period, period_index, table, payments)

        """
        offset - это индекс, по которому находится строка-шаблон платежа. Он вычисляется
        следующим образом. 4 - это первые четыре строки в таблице с контрактом. period_index - 
        это индекс текущего блока с периодом, а 3 - это 3 фиксированных строки в каждом таком 
        блоке (Начислено за период, корректировка обязательств и итог). payments - это число 
        платежей, которые уже заполнены в текущей таблице. Последняя 1 - это фиксированная 
        строка Начислено за период. 
        В итоге в таблицу по индексу offset + 1 вставляются новые строки с платежами, которые нужно
        заполнить
        """
        offset = 4 + period_index * 3 + payments + 1
        self.clone_payment_row(table, len(period["payments_2"]) - 1, offset + 2, offset + 3)
        self.clone_payment_row(table, len(period["payments_1"]) - 1, offset, offset + 1)

        for i, payment in enumerate(period["payments_1"]):
            self.fill_payment_1(payment, i, table, offset)

        for i, payment in enumerate(period["payments_2"]):
            self.fill_payment_2(payment, i, table, offset + len(period["payments_1"]) + 1)


    def fill_common_period_info(self, period, period_index: int, table: Table, payments: int):
        to_replace = (
            (self.borders("период месяц.год"),              period["period"]),
            (self.borders("общая сумма начисления"),        period["first_debt"]),
            (self.borders("дата начала начисления"),        period["first_date"]),
            (self.borders("общая сумма корректировки"),     period["second_debt"]),
            (self.borders("дата начала корректировки"),     period["second_date"]),
            (self.borders("сумма пени за период"),          period["result"]),
        )

        stride = 3
        cell = table.row_cells(4 + period_index * stride + payments)[0]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])

        cell = table.row_cells(4 + period_index * stride + payments)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[1][0], to_replace[1][1])
        cell = table.row_cells(4 + period_index * stride + payments)[3]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[2][0], to_replace[2][1])
        cell = table.row_cells(4 + period_index * stride + payments)[4]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])

        cell = table.row_cells(6 + period_index * stride + payments)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[3][0], to_replace[3][1])
        cell = table.row_cells(6 + period_index * stride + payments)[3]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[4][0], to_replace[4][1])

        cell = table.row_cells(8 + period_index * stride + payments)[11]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[5][0], to_replace[5][1])


    def fill_payment_1(self, payment, payment_index: int, table: Table, offset: int):
        to_replace = (
            (self.borders("сумма платежа"),                 payment["debt"]),
            (self.borders("дата начала просрочки"),         payment["start_date"]),
            (self.borders("дата конца просрочки"),          payment["end_date"]),
            (self.borders("кол-во дней просрочки"),         payment["period_days"]),
            (self.borders("ставка"),                        payment["interest_rate"]),
            (self.borders("доля ставки"),                   payment["share"]),
            (self.borders("формула"),                       payment["formulae"]),
            (self.borders("пени"),                          payment["peny"]),
        )
        cell = table.row_cells(offset + payment_index)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])
        cell = table.row_cells(offset + payment_index)[3]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[1][0], to_replace[1][1])
        cell = table.row_cells(offset + payment_index)[4]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[2][0], to_replace[2][1])
        cell = table.row_cells(offset + payment_index)[5]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[3][0], to_replace[3][1])
        
        cell = table.row_cells(offset + payment_index)[6]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[4][0], to_replace[4][1])
        cell = table.row_cells(offset + payment_index)[8]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[5][0], to_replace[5][1])
        cell = table.row_cells(offset + payment_index)[9]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[6][0], to_replace[6][1])
        cell = table.row_cells(offset + payment_index)[11]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[7][0], to_replace[7][1])


    def fill_payment_2(self, payment, payment_index: int, table: Table, offset: int):
        to_replace = (
            (self.borders("сумма корректировки"),                 payment["debt"]),
            (self.borders("дата начала корректировки"),         payment["start_date"]),
            (self.borders("дата конца корректировки"),          payment["end_date"]),
            (self.borders("кол-во дней корректировки"),         payment["period_days"]),
            (self.borders("ставка"),                        payment["interest_rate"]),
            (self.borders("доля ставки"),                   payment["share"]),
            (self.borders("формула"),                       payment["formulae"]),
            (self.borders("пени"),                          payment["peny"]),
        )
        cell = table.row_cells(offset + payment_index)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])
        cell = table.row_cells(offset + payment_index)[3]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[1][0], to_replace[1][1])
        cell = table.row_cells(offset + payment_index)[4]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[2][0], to_replace[2][1])
        cell = table.row_cells(offset + payment_index)[5]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[3][0], to_replace[3][1])
        
        cell = table.row_cells(offset + payment_index)[6]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[4][0], to_replace[4][1])
        cell = table.row_cells(offset + payment_index)[8]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[5][0], to_replace[5][1])
        cell = table.row_cells(offset + payment_index)[9]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[6][0], to_replace[6][1])
        cell = table.row_cells(offset + payment_index)[11]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[7][0], to_replace[7][1])


    def borders(self, string: str, start: str = "/*", end: str = "*/"):
        return start + string + end


    def fill_first_table(self):
        text_1 = self.redactor.get_table(0).row_cells(4)[4].text
        text_2 = self.redactor.get_table(0).row_cells(6)[4].text

        # self.clone_row_in_first_table_2(0, 4, 5, text_1)
        self.clone_row_in_first_table_1(0, 5, 6)
        # self.clone_row_in_first_table_1(0, 9, 10)

        # self.clone_row_in_first_table_2(0, 6, 11, text_2)


    def clone_row_in_first_table_1(self, table_index, source_row_index, target_row_index):
        table = self.redactor.get_table(table_index)
        c1 = table.row_cells(source_row_index)[0]
        new_row = self.redactor.insert_row_in_table(table, target_row_index)
        self.redactor.clone_table_row(table.rows[source_row_index], new_row)
        c2 = new_row.cells[0]
        self.redactor.merge_table_cells(c1, c2)

        self.redactor.merge_table_cells(new_row.cells[1], new_row.cells[2])
        self.redactor.merge_table_cells(new_row.cells[6], new_row.cells[7])
        self.redactor.merge_table_cells(new_row.cells[9], new_row.cells[10])


    def clone_row_in_first_table_2(self, table_index, source_row_index, target_row_index, text):
        table = self.redactor.get_table(table_index)
        c1 = table.row_cells(source_row_index)[0]
        new_row = self.redactor.insert_row_in_table(table, target_row_index)
        self.redactor.clone_table_row(table.rows[source_row_index], new_row)
        c2 = new_row.cells[0]
        self.redactor.merge_table_cells(c1, c2)

        self.redactor.merge_table_cells(new_row.cells[1], new_row.cells[2])
        for i in range(5, 12):
            self.redactor.merge_table_cells(new_row.cells[4], new_row.cells[i])

        new_row.cells[4].paragraphs[0].runs[0].text = text.strip()


    def fill_second_table(self):
        table = self.redactor.get_table(1)
        new_row = self.redactor.insert_row_in_table(table, 2)
        self.redactor.clone_table_row(table.rows[1], new_row)


    def clone_table(self, count: int = 1):
        copied_table = self.redactor.get_table(0)

        for i in range(count):
            new_table = self.redactor.insert_table_after_table(copied_table)
            new_paragraph = self.redactor.insert_paragraph_after_table(copied_table)
            self.redactor.clone_table(copied_table, new_table)

            # Объединяем ячейки первой строки
            self.merge_row_1(new_table, 0)

            # Объединяем ячейки второй строки
            self.merge_row_2(new_table, 1)

            # Объединяем ячейки третьей и четвертой строк
            self.merge_row_3_4(new_table, 2)

            # Объединяем ячейки пятой и седьмой строк
            self.merge_row_5(new_table, 4)
            self.merge_row_5(new_table, 6, True)

            # Объединяем ячейки шестой, восьмой и девятой строк
            self.merge_row_6(new_table, 5)
            self.merge_row_6(new_table, 7)
            self.merge_row_6(new_table, 8)

            # Объединяем ячейки десятой и одиннадцатой строк
            self.merge_row_10(new_table, 9)
            self.merge_row_10(new_table, 10)


    def merge_row_1(self, table, row_index):
        # Объединяем ячейки первой строки
        row = table.rows[row_index]
        self.redactor.merge_table_cells(row.cells[0], row.cells[1])
        for i in range(3, 12):
            self.redactor.merge_table_cells(row.cells[2], row.cells[i])


    def merge_row_2(self, table, row_index):
        # Объединяем ячейки второй строки
        row = table.rows[row_index]
        self.redactor.merge_table_cells(row.cells[0], row.cells[1])
        for i in range(3, 7):
            self.redactor.merge_table_cells(row.cells[2], row.cells[i])
        for i in range(8, 10):
            self.redactor.merge_table_cells(row.cells[7], row.cells[i])
        self.redactor.merge_table_cells(row.cells[10], row.cells[11])


    def merge_row_3_4(self, table, row_3_index):
        # Объединяем ячейки третьей и четвертой строк
        row_3 = table.rows[row_3_index]
        row_4 = table.rows[row_3_index + 1]
        for row in [row_3, row_4]:
            self.redactor.merge_table_cells(row.cells[1], row.cells[2])
            self.redactor.merge_table_cells(row.cells[6], row.cells[7])
            self.redactor.merge_table_cells(row.cells[9], row.cells[10])

        for i in range(4, 6):
            self.redactor.merge_table_cells(row_3.cells[3], row_3.cells[i])

        for i in [0, 1, 6, 7, 8, 9, 10, 11]:
            self.redactor.merge_table_cells(row_3.cells[i], row_4.cells[i])


    def merge_row_5(self, table, row_index, first_cell_flag = False):
        """
        first_cell_flag - флаг сообщает, нужно ли присоединять первую колонку текущей строки
        к общей колонке "период". По умолчанию не нужно.
        """
        row = table.rows[row_index]
        self.redactor.merge_table_cells(row.cells[1], row.cells[2])
        for i in range(5, 12):
            self.redactor.merge_table_cells(row.cells[4], row.cells[i])

        if first_cell_flag:
            previous_row = table.rows[row_index - 1]
            self.redactor.merge_table_cells(previous_row.cells[0], row.cells[0])


    def merge_row_6(self, table, row_index):
        row = table.rows[row_index]
        self.redactor.merge_table_cells(row.cells[1], row.cells[2])
        self.redactor.merge_table_cells(row.cells[6], row.cells[7])
        self.redactor.merge_table_cells(row.cells[9], row.cells[10])

        previous_row = table.rows[row_index - 1]
        self.redactor.merge_table_cells(previous_row.cells[0], row.cells[0])


    def merge_row_10(self, table, row_index):
        row = table.rows[row_index]
        for i in range(1, 12):
            self.redactor.merge_table_cells(row.cells[0], row.cells[i])


    def clone_block(self, table: Table, index: int, count: int = 1):
        """Клонирует блок. 1 блок - 1 период

        Args:
            index (int): индекс первой строки копируемого блока
            count (int): количество копий. Все копии будут сделаны
        """

        stride = 5
        for i in range(count):
            for j in range(stride):
                new_row = self.redactor.insert_row_in_table(table, index + (i + 1) * stride + j)
                self.redactor.clone_table_row(table.rows[index + j], new_row)

            # self.merge_row_5(table.rows[index + (i + 1) * stride])
            self.merge_row_5(table, index + (i + 1) * stride)
            # self.merge_row_6(table.rows[index + (i + 1) * stride + 1])
            self.merge_row_6(table, index + (i + 1) * stride + 1)
            # self.merge_row_5(table.rows[index + (i + 1) * stride + 2])
            self.merge_row_5(table, index + (i + 1) * stride + 2, True)
            # self.merge_row_6(table.rows[index + (i + 1) * stride + 3])
            self.merge_row_6(table, index + (i + 1) * stride + 3)
            # self.merge_row_6(table.rows[index + (i + 1) * stride + 4])
            self.merge_row_6(table, index + (i + 1) * stride + 4)
            # self.merge_cols_period(table, index + (i + 1) * stride)


    def clone_payment_row(self, table: Table, count: int = 1, source_row_index: int = 5, target_row_index: int = 6):
        """Клонирует строку с данными платежа.

        Args:
            table (Table): Таблица, в которой происходят изменения
            source_row_index (int, optional): Индекс строки, которую копируем (она является шаблонной). 
            По умолчанию = 4.
            target_row_index (int, optional): Индекс, по которому вставляем копию строки. По 
            умолчанию вставляется сразу после шаблонной строки.
        """
        for i in range(count):
            new_row = self.redactor.insert_row_in_table(table, target_row_index)
            self.redactor.clone_table_row(table.rows[source_row_index], new_row)
            self.merge_row_6(table, target_row_index)
            self.redactor.merge_table_cells(table.rows[target_row_index - 1].cells[0], table.rows[target_row_index].cells[0])
