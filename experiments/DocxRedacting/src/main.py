from docx import Document

from DocTemplateReplacer import DocTemplateReplacer

from ConfigParsing.ConfigParser import parse_config
from DocEditing.CloneTemplate import clone_template
from DocEditing.EditParagraphs import edit_paragraphs
from DocEditing.EditTables import edit_tables


def old_create_document(config_filename: str, template_filename: str, output_filename: str = "output.docx") -> None:
    # 1. Загружаем конфиг
    config = parse_config(config_filename)

    # 2. Копируем шаблон документа
    clone_template(template_filename, output_filename)
    doc = Document(output_filename)

    # 3. Меняем текст
    edit_paragraphs(doc, config)
    edit_tables(doc, config)

    doc.save(output_filename)


def create_document(config_filename: str, template_filename: str, output_filename: str = "output.docx") -> None:
    replacer = DocTemplateReplacer()
    replacer.make_instance(config_filename, template_filename, output_filename)


if __name__ == "__main__":
    config_filename = "lawsuit_create.json"
    template_filename = "template.docx"
    output_filename = "output.docx"
    create_document(config_filename, template_filename, output_filename)

    # doc = Document(template_filename)
    # text = []
    # for n, paragraph in enumerate(doc.paragraphs):
    #     text.append(f"[{n}]: " + paragraph.text)
    # print('\n\n'.join(text))



"""

2 абзац - /*Цена иска*/
3 абзац - /*Госпошлина*/

9 абзац - /*ответчик*/

11 абзац - /*тип договора*/
12 абзац - /*тип договора*/
13 абзац - /*сумма долга*/

15 абзац (дублируется) - /*номер раздела*/, /*пункт*/, /*номер договора*/

24 абзац (дублируется) - /*номер претензии*/

29 абзац - после него идёт таблица, которую надо редактировать

31 абзац - /*цена иска*/
32 абзац - /*госпошлина*/

35 абзац - /*ответчик*/, /*цена иска*/

36-39 абзацы дублируются
37 абзац - /*Задолженность*/, /*период*/
38 абзац - /*неустойка*/

40 абзац - /*ответчик*/, /*огрн*/, /*инн*/, /*госпошлина*/


"""


