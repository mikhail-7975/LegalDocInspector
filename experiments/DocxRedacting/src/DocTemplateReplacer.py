import json

from docx import Document

from DocxRedactor import DocxRedactor


class DocTemplateReplacer:
    """
    Класс создания документа по шаблону. При создании экземпляра принимает три аргумента:
    
    """

    def __init__(self) -> None:
        self.redactor: DocxRedactor = DocxRedactor()
        self.config: dict = None
        self.doc: Document = None


    def make_instance(self, config_filename: str, template_filename: str, output_filename: str = None):
        """
        config_filename: str - имя файла с конфигами для шаблона
        template_filename: str - имя файла-шаблона
        output_filename: str - имя файла, созданного по шаблону
        """
        self.config = self.parse_config(config_filename)

        self.redactor.clone_file(template_filename, output_filename)
        self.doc = self.redactor.open(output_filename)

        self.fill_file()

        self.redactor.save()
        self.redactor.close()


    def parse_config(self, filename: str) -> dict:
        with open(filename, 'r', encoding="utf-8") as file:
            data = json.load(file)
        return data


    def borders(self, string: str, start: str = "/*", end: str = "*/"):
        return start + string + end


    def fill_file(self):
        self.fill_first_table()
        self.fill_second_table()
        self.fill_third_table()
        self.fill_first_list()
        self.fill_second_list()
        self.fill_third_list()
        self.fill_other_parts()


    def fill_first_table(self):
        to_replace = (
            (self.borders("ответчик"),    "#ответчик#".upper()),
            (self.borders("огрн"),        self.config["plaintiff_info"]["ogrn"],),
            (self.borders("инн"),         self.config["plaintiff_info"]["inn"])
        )

        table = self.redactor.get_table(0)
        cell = table.row_cells(6)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])
        self.redactor.replace_text_in_paragraph(cell.paragraphs[1], to_replace[1][0], to_replace[1][1])
        self.redactor.replace_text_in_paragraph(cell.paragraphs[1], to_replace[2][0], to_replace[2][1])


    def fill_second_table(self):
        rows_n = len(self.config["contracts_info"]) # Кол-во элементов в списке договоров
        table = self.redactor.get_table(1)

        for i in range(rows_n):
            if i < rows_n - 1:
                new_row = self.redactor.insert_row_in_table(table, 2 + i)
                self.redactor.clone_table_row(table.rows[1 + i], new_row)
                # for cell in table.row_cells(1 + i):
                #     # self.redactor._set_borders_to_cell(cell)
                #     pass
                #     # self.redactor.set_vertical_alignment_to_cell(cell, "center")

            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[0].paragraphs[0],
                self.borders("номер договора"),
                self.config["contracts_info"][i][1]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[1].paragraphs[0],
                self.borders("период"),
                self.config["contracts_info"][i][2]["contract_periods"]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[2].paragraphs[0],
                self.borders("задолженность"),
                self.config["contracts_info"][i][2]["debt"]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[3].paragraphs[0],
                self.borders("срок оплаты"),
                "#срок оплаты#".upper()
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[4].paragraphs[0],
                self.borders("пункт"),
                self.config["contracts_info"][i][2]["contract_point"]
            )


    def fill_third_table(self):
        rows_n = len(self.config["contracts_info"]) # Кол-во элементов в списке договоров
        table = self.redactor.get_table(2)

        for i in range(rows_n):
            if i < rows_n - 1:
                new_row = self.redactor.insert_row_in_table(table, 2 + i)
                self.redactor.clone_table_row(table.rows[1 + i], new_row)
                # for cell in table.row_cells(1 + i):
                #     self.redactor._set_borders_to_cell(cell)
                #     self.redactor.set_vertical_alignment_to_cell(cell, "center")

            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[0].paragraphs[0],
                self.borders("номер договора"),
                self.config["contracts_info"][i][1]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[1].paragraphs[0],
                self.borders("период"),
                self.config["contracts_info"][i][2]["contract_periods"]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[2].paragraphs[0],
                self.borders("задолженность"),
                self.config["contracts_info"][i][2]["debt"]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[3].paragraphs[0],
                self.borders("неустойка"),
                "#неустойка#".upper()
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[4].paragraphs[0],
                self.borders("неустойка+задолженность"),
                "#неустойка+задолженность#".upper()
            )


    def fill_first_list(self):
        rows_n = len(self.config["contracts_info"]) # Кол-во элементов в списке договоров
        template_text = f"раздел {self.borders("номер раздела")}, в том числе пункт {self.borders("пункт")} договора № {self.borders("номер договора")}"
        start = self.redactor.find_paragraph_consists_of_text(template_text)   # Индекс первого абзаца списка

        for i in range(rows_n):
            if i < rows_n - 1:
                src_p = self.redactor.get_paragraph(start + i)
                new_p = self.redactor.insert_paragraph(start + i + 1)
                self.redactor.clone_paragraph(src_p, new_p)

            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i),
                self.borders("номер раздела"),
                "#номер раздела#".upper()
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i),
                self.borders("пункт"),
                self.config["contracts_info"][i][2]["contract_point"]
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i),
                self.borders("номер договора"),
                self.config["contracts_info"][i][1]
            )


    def fill_second_list(self):
        rows_n = len(self.config["lawsuit_info"]["claims"]) # Кол-во элементов в списке претензий
        template_text = f"list{self.borders("номер претензии")};"
        start = self.redactor.find_paragraph_consists_of_text(template_text)   # Индекс первого абзаца списка

        for i in range(rows_n):
            if i < rows_n - 1:
                src_p = self.redactor.get_paragraph(start + i)
                new_p = self.redactor.insert_paragraph(start + i + 1)
                self.redactor.clone_paragraph(src_p, new_p)

            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i),
                template_text[:len(template_text) - 1],
                self.config["lawsuit_info"]["claims"][i]
            )


    def fill_third_list(self):
        rows_n = len(self.config["contracts_info"]) # Кол-во элементов в списке договоров
        template_text = "по Договору № 05.403297-ТЭ от 01.08.2017:"
        start = self.redactor.find_paragraph_consists_of_text(template_text)   # Индекс первого абзаца списка

        stride = 4
        for i in range(rows_n):
            if i < rows_n - 1:
                for j in range(stride):
                    src_p = self.redactor.get_paragraph(start + i * stride + j)
                    new_p = self.redactor.insert_paragraph(start + (i + 1) * stride + j)
                    self.redactor.clone_paragraph(src_p, new_p)

            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride),
                "05.403297-ТЭ от 01.08.2017",
                self.config["contracts_info"][i][1]
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride + 1),
                self.borders("задолженность"),
                self.config["contracts_info"][i][2]["debt"]
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride + 1),
                self.borders("период"),
                self.config["contracts_info"][i][2]["contract_periods"]
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride + 2),
                self.borders("неустойка"),
                self.config["contracts_info"][i][2]["debt_penalty"]
            )


    def fill_other_parts(self):
        to_replace = {
            "/*цена иска*/": self.config["lawsuit_info"]["cost"],
            "/*госпошлина*/": self.config["lawsuit_info"]["tax"],
            "/*ответчик*/": "/*ОТВЕТЧИК*/",
            "/*тип договора*/": self.config["lawsuit_info"]["service_type"],
            "/*сумма долга*/": "/*СУММА ДОЛГА*/",
            "/*огрн*/": self.config["plaintiff_info"]["ogrn"],
            "/*инн*/": self.config["plaintiff_info"]["inn"]
        }

        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                for k, v in to_replace.items():
                    if k in run.text:
                        run.text = run.text.replace(k, v)
