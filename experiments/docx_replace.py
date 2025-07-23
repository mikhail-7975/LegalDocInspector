from docx import Document

def replace_text_preserve_formatting(docx_path, replaces, output_path):
    """
    Replace text in a DOCX document while preserving formatting.

    Args:
        docx_path (str): Path to the input DOCX file
        old_text (str): Text to be replaced
        new_text (str): New text to insert
        output_path (str): Path to save the modified document
    """
    doc = Document(docx_path)

    for r in replaces:
        old_text, new_text = r
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                inline = paragraph.runs
                for i in range(len(inline)):
                    if old_text in inline[i].text:
                        text = inline[i].text.replace(old_text, new_text)
                        inline[i].text = text

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if old_text in paragraph.text:
                            inline = paragraph.runs
                            for i in range(len(inline)):
                                if old_text in inline[i].text:
                                    text = inline[i].text.replace(old_text, new_text)
                                    inline[i].text = text

    doc.save(output_path)

# Example usage
input_file = "input.docx"
output_file = "generated_document.docx"
replace_text_preserve_formatting(input_file,
                                 [
                                     ("Московская объединенная энергетическая компания", "Это документ для проверки форматирования"),
                                     ("ИНН 7720518494", "ИНН 88888888"),
                                     ("ДЕЗ-СТОЛИЦА", "РОМАШКА"),
                                     ("05.403297-ТЭ", "12.345678-ТЭ")
                                 ], output_file)