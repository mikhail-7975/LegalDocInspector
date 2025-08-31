import json
from datetime import datetime, timedelta

from docx import Document

from DocxRedactor import DocxRedactor



class ClaimGenerator:
    """
    Класс создания документа по шаблону. При создании экземпляра принимает три аргумента:
    
    """

    def __init__(self) -> None:
        self.redactor: DocxRedactor = DocxRedactor()
        self.config: dict = None
        self.doc: Document = None


    def make_instance(self, config: str, template_filename: str, output_filename: str = None):
        """
        config_filename: str - конфиги для шаблона
        template_filename: str - имя файла-шаблона
        output_filename: str - имя файла, созданного по шаблону
        """
        self.config = config

        self.redactor.clone_file(template_filename, output_filename)
        self.doc = self.redactor.open(output_filename)

        self.fill_file()

        self.redactor.save()
        self.redactor.close()


    def parse_config(self, filename: str) -> dict:
        with open(filename, 'r', encoding="utf-8") as file:
            data = json.load(file)
        return data


    def borders(self, string: str, start: str = "/*", end: str = "*/"):
        return start + string + end


    def fill_file(self):
        # self.redactor.print_table(self.redactor.get_table(0))
        self.fill_first_table()
        self.fill_second_table()
        self.fill_third_table()
        self.fill_first_list()
        self.fill_second_list()
        self.fill_third_list()
        self.fill_other_parts()


    def fill_first_table(self):
        # self.redactor.print_table(table)

        to_replace = (
            # строка 2
            (self.borders("истец полное имя"),      self.config["plaintiff_info"]["full_name"]),
            (self.borders("истец огрн"),            self.config["plaintiff_info"]["ogrn"]),
            (self.borders("истец инн"),             self.config["plaintiff_info"]["inn"]),

            # Строка 3
            (self.borders("истец адрес"),           self.config["plaintiff_info"]["addres"]),

            # Строка 6
            (self.borders("ответчик полное имя"),   self.config["defendant_info"]["full_name"]),
            (self.borders("ответчик огрн"),         self.config["defendant_info"]["ogrn"]),
            (self.borders("ответчик инн"),          self.config["defendant_info"]["inn"]),

            # Строка 7
            (self.borders("ответчик адрес"),        self.config["defendant_info"]["addres"]),
        )

        table = self.redactor.get_table(0)

        cell = table.row_cells(2)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])
        self.redactor.replace_text_in_paragraph(cell.paragraphs[1], to_replace[1][0], to_replace[1][1])
        self.redactor.replace_text_in_paragraph(cell.paragraphs[1], to_replace[2][0], to_replace[2][1])

        cell = table.row_cells(3)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[3][0], to_replace[3][1])

        cell = table.row_cells(6)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[4][0], to_replace[4][1])
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[5][0], to_replace[5][1])
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[6][0], to_replace[6][1])

        cell = table.row_cells(7)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[7][0], to_replace[7][1])

        self.redactor.save()


    def fill_second_table(self):
        rows_n = len(self.config["contracts_info"]) # Кол-во элементов в списке договоров
        table = self.redactor.get_table(1)

        for i in range(rows_n):
            if i < rows_n - 1:
                new_row = self.redactor.insert_row_in_table(table, 2 + i)
                self.redactor.clone_table_row(table.rows[1 + i], new_row)
                # for cell in table.row_cells(1 + i):
                #     # self.redactor._set_borders_to_cell(cell)
                #     pass
                #     # self.redactor.set_vertical_alignment_to_cell(cell, "center")

            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[0].paragraphs[0],
                self.borders("номер договора"),
                self.config["contracts_info"][i][1]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[1].paragraphs[0],
                self.borders("период"),
                self.config["contracts_info"][i][2]["contract_periods"]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[2].paragraphs[0],
                self.borders("задолженность"),
                self.config["contracts_info"][i][2]["debt"]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[3].paragraphs[0],
                self.borders("срок оплаты"),
                # "#срок оплаты#".upper()
                self.config["contracts_info"][i][2]["last_day"]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[4].paragraphs[0],
                self.borders("пункт"),
                self.config["contracts_info"][i][2]["contract_point"]
            )

        self.redactor.replace_text_in_paragraph(
            table.row_cells(rows_n + 1)[2].paragraphs[0],
            self.borders("сумма долга"),
            self.config["table_info"]["all_debt"]
        )


    def fill_third_table(self):
        rows_n = len(self.config["contracts_info"]) # Кол-во элементов в списке договоров
        table = self.redactor.get_table(2)

        for i in range(rows_n):
            if i < rows_n - 1:
                new_row = self.redactor.insert_row_in_table(table, 2 + i)
                self.redactor.clone_table_row(table.rows[1 + i], new_row)
                # for cell in table.row_cells(1 + i):
                #     self.redactor._set_borders_to_cell(cell)
                #     self.redactor.set_vertical_alignment_to_cell(cell, "center")

            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[0].paragraphs[0],
                self.borders("номер договора"),
                self.config["contracts_info"][i][1]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[1].paragraphs[0],
                self.borders("период"),
                self.config["contracts_info"][i][2]["contract_periods"]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[2].paragraphs[0],
                self.borders("задолженность"),
                self.config["contracts_info"][i][2]["debt"]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[3].paragraphs[0],
                self.borders("неустойка"),
                # "#неустойка#".upper()
                self.config["contracts_info"][i][2]["penalty"]
            )
            self.redactor.replace_text_in_paragraph(
                table.row_cells(i + 1)[4].paragraphs[0],
                self.borders("неустойка+задолженность"),
                # "#неустойка+задолженность#".upper()
                self.config["contracts_info"][i][2]["debt_penalty"]
            )

        self.redactor.replace_text_in_paragraph(
            table.row_cells(rows_n + 1)[2].paragraphs[0],
            self.borders("сумма долга"),
            self.config["table_info"]["all_debt"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(rows_n + 1)[3].paragraphs[0],
            self.borders("неустойка общая"),
            self.config["table_info"]["all_penalty"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(rows_n + 1)[4].paragraphs[0],
            self.borders("цена иска"),
            self.config["table_info"]["cost_of_lawsuit"]
        )


    def fill_first_list(self):
        rows_n = len(self.config["contracts_info"]) # Кол-во элементов в списке договоров
        template_text = f"раздел {self.borders("номер раздела")}, в том числе пункт {self.borders("пункт")} договора № {self.borders("номер договора")}"
        start = self.redactor.find_paragraph_consists_of_text(template_text)   # Индекс первого абзаца списка

        for i in range(rows_n):
            if i < rows_n - 1:
                src_p = self.redactor.get_paragraph(start + i)
                new_p = self.redactor.insert_paragraph(start + i + 1)
                self.redactor.clone_paragraph(src_p, new_p)

            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i),
                self.borders("номер раздела"),
                # "#номер раздела#".upper()
                self.config["contracts_info"][i][2]["contract_point"].split(".")[0]
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i),
                self.borders("пункт"),
                self.config["contracts_info"][i][2]["contract_point"]
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i),
                self.borders("номер договора"),
                self.config["contracts_info"][i][1]
            )


    def fill_second_list(self):
        rows_n = len(self.config["lawsuit_info"]["claims"]) # Кол-во элементов в списке претензий
        template_text = f"list{self.borders("номер претензии")};"
        start = self.redactor.find_paragraph_consists_of_text(template_text)   # Индекс первого абзаца списка

        for i in range(rows_n):
            if i < rows_n - 1:
                src_p = self.redactor.get_paragraph(start + i)
                new_p = self.redactor.insert_paragraph(start + i + 1)
                self.redactor.clone_paragraph(src_p, new_p)

            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i),
                template_text[:len(template_text) - 1],
                self.config["lawsuit_info"]["claims"][i]
            )


    def fill_third_list(self):
        rows_n = len(self.config["contracts_info"]) # Кол-во элементов в списке договоров
        template_text = "по Договору № 05.403297-ТЭ от 01.08.2017:"
        start = self.redactor.find_paragraph_consists_of_text(template_text)   # Индекс первого абзаца списка

        # Список индексов тех параграфов, которые нужно удалить после цикла. Почему их нужно удалять?
        # Эти параграфы - это корректировка обязательств, которая есть не у всех договоров. Поэтому 
        # там где ее нет, параграфы удаляются
        delete_par_indices = []

        # Эта переменная нужна для подсчета числа заполненных строк в списке. Нельзя просто сделать
        # rows_n * stride, поскольку у некоторых договоров отсутствует годовая корректировка, и 
        # соответствующие параграфы удаляются из списка, т.е. строк становится меньше
        filled_rows = 0

        stride = 5
        for i in range(rows_n):
            if i < rows_n - 1:
                for j in range(stride):
                    src_p = self.redactor.get_paragraph(start + i * stride + j)
                    new_p = self.redactor.insert_paragraph(start + (i + 1) * stride + j)
                    self.redactor.clone_paragraph(src_p, new_p)

            contract_number = list(self.config["table_info"].keys())[i]

            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride),
                "05.403297-ТЭ от 01.08.2017",
                # self.config["contracts_info"][i][1]
                contract_number
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride + 1),
                self.borders("задолженность"),
                # self.config["contracts_info"][i][2]["debt"]
                self.config["table_info"][contract_number]["debt"]
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride + 1),
                self.borders("период"),
                # self.config["contracts_info"][i][2]["contract_periods"]
                self.config["table_info"][contract_number]["contract_periods"]
            )

            if self.config["table_info"][contract_number]["contract_periods_correcting"] is not None:
                if self.config["table_info"][contract_number]["correcting_debt"] != "0,00":
                    self.redactor.replace_text_in_paragraph(
                        self.redactor.get_paragraph(start + i * stride + 2),
                        self.borders("доля годовой корректировки"),
                        # self.config["contracts_info"][i][2]["debt_penalty"]
                        self.config["table_info"][contract_number]["correcting_debt"]
                    )
                    self.redactor.replace_text_in_paragraph(
                        self.redactor.get_paragraph(start + i * stride + 2),
                        self.borders("период корректировки"),
                        # self.config["contracts_info"][i][2]["debt_penalty"]
                        self.config["table_info"][contract_number]["contract_periods_correcting"]
                    )
                else:
                    if start + i * stride + 2 not in delete_par_indices:
                        delete_par_indices.append(start + i * stride + 2)
            else:
                if start + i * stride + 2 not in delete_par_indices:
                    delete_par_indices.append(start + i * stride + 2)

            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride + 3),
                self.borders("неустойка"),
                # self.config["contracts_info"][i][2]["debt_penalty"]
                self.config["table_info"][contract_number]["penalty"]
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride + 3),
                self.borders("период неустойки 1"),
                # self.config["contracts_info"][i][2]["debt_penalty"]
                self.config["table_info"][contract_number]["penalty_period"].split(" по ")[0]
            )
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride + 3),
                self.borders("период неустойки 2"),
                # self.config["contracts_info"][i][2]["debt_penalty"]
                self.config["table_info"][contract_number]["penalty_period"].split(" по ")[1]
            )

            date_string = self.config["table_info"][contract_number]["penalty_period"].split(" по ")[1]
            format_string = "%d.%m.%Y"
            parsed_date = datetime.strptime(date_string, format_string)
            parsed_date_plus_1 = parsed_date + timedelta(days=1)
            self.redactor.replace_text_in_paragraph(
                self.redactor.get_paragraph(start + i * stride + 4),
                self.borders("начальная дата неустойки"),
                # self.config["contracts_info"][i][2]["debt_penalty"]
                parsed_date_plus_1.strftime(format_string)
            )

        delete_paragraps = 0
        for index in delete_par_indices:
            self.redactor.delete_paragraph(index - delete_paragraps)
            delete_paragraps += 1


    def fill_other_parts(self):
        to_replace = {
            "/*цена иска*/": self.config["lawsuit_info"]["cost"],
            "/*госпошлина*/": self.config["lawsuit_info"]["tax"],
            "/*истец полное имя*/": self.config["plaintiff_info"]["full_name"],
            "/*истец сокращенное имя*/": self.config["plaintiff_info"]["short_name"],
            "/*ответчик полное имя*/": self.config["defendant_info"]["full_name"],
            "/*ответчик сокращенное имя*/": self.config["defendant_info"]["short_name"],
            "/*тип договора*/": self.config["lawsuit_info"]["service_type"],
            "/*сумма долга*/": self.config["table_info"]["all_debt"],
            "/*истец огрн*/": self.config["plaintiff_info"]["ogrn"],
            "/*истец инн*/": self.config["plaintiff_info"]["inn"],
            "/*ответчик огрн*/": self.config["defendant_info"]["ogrn"],
            "/*ответчик инн*/": self.config["defendant_info"]["inn"],
        }

        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                for k, v in to_replace.items():
                    if k in run.text:
                        run.text = run.text.replace(k, v)
