import json
from copy import deepcopy

from docx import Document
from docx.table import Table

from DocxRedactor import DocxRedactor
from datetime import datetime


class CalculationClaimGenerator:
    """
    Класс для создания документа с расчётом для иска
    
    """
    def __init__(self) -> None:
        self.redactor: DocxRedactor = DocxRedactor()
        self.config: dict = None
        self.config2: dict = None
        self.doc: Document = None


    def make_instance(self, config: str, config2: str, template_filename: str, output_filename: str = None):
        """
        config: str - конфиги для шаблона
        config2: str - конфиги для нижней таблицы
        template_filename: str - имя файла-шаблона
        output_filename: str - имя файла, созданного по шаблону
        """
        self.config = self.convert_data_from_calculator(config)
        self.config2 = config2

        self.redactor.clone_file(template_filename, output_filename)
        self.doc = self.redactor.open(output_filename)

        self.fill_second_table()
        self.fill_other_parts()
        self.fill_file()

        self.redactor.save()
        self.redactor.close()


    def convert_data_from_calculator(self, contracts):
        """На вход метод принимает данные из калькулятора (списком, т.е. в contracts может находится несколько
        наборов данных полученных от калькулятора), каждый элемент списка - это данные об одном контракте.
        
        На выходе возвращаются те же самые данные, но в структуре, удобной для записи в расчёт к иску.
        """

        with open("temp.json", "w") as file:
            json.dump(contracts, file, ensure_ascii=False, indent=4)

        converted_contracts = {"contracts": []}

        for contract in contracts:
            converted_contract = {
                "contract_number": contract["contract_number"],
                "start_date_of_delay": contract["start_of_table"]["start"],
                "end_date_of_delay": contract["start_of_table"]["end"],
                "total_debt": contract["end_of_table1"]["money"],
                "total_peny": contract["end_of_table2"]["money"],
                "periods": [],
            }

            for period in contract.keys():
                if period not in ["contract_number", "start_of_table", "end_of_table1", "end_of_table2", "debt_info"]:
                    new_period = {}
                    new_period["period"] = period

                    for item in contract[period]:

                        if item["text"] == "Итого:":
                            new_period["total"] = item["penalty"]

                    new_period["rows"] = []
                    for item in contract[period]:

                        if "type" in item.keys():
                            if item["type"] == "debt_info":
                                continue

                        if item["text"] != "Итого:":
                            new_period["rows"].append(item)

                    converted_contract["periods"].append(new_period)

            converted_contracts["contracts"].append(converted_contract)

        return converted_contracts


    def fill_file(self):
        self.clone_table(len(self.config["contracts"]) - 1)

        # self.redactor.print_table(self.redactor.get_table(0))

        for i, contract in enumerate(self.config["contracts"]):
            self.fill_contract(contract, i)


    def fill_contract(self, contract, contract_index):
        table = self.redactor.get_table(contract_index)

        self.fill_common_contract_info(contract, table)

        self.clone_block(table, 4, len(contract["periods"]) - 1)

        filled_rows = 0     # число заполненных строк во всей таблице
        for i, period in enumerate(contract["periods"]):
            res = self.fill_period(period, i, table, filled_rows)
            filled_rows += res
            # filled_rows += 3
            # payments += len(period["payments_1"])
        #     payments += len(period["payments_2"])


    def fill_common_contract_info(self, contract, table: Table):
        to_replace = (
            (self.borders("номер договора"),            contract["contract_number"]),
            (self.borders("дата начала просрочки"),     contract["start_date_of_delay"]),
            (self.borders("дата конца просрочки"),      contract["end_date_of_delay"]),
            (self.borders("сумма долга"),               contract["total_debt"]),
            (self.borders("сумма пеней"),               contract["total_peny"]),
        )

        cell = table.row_cells(0)[2]
        flag = self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])
        cell = table.row_cells(1)[2]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[1][0], to_replace[1][1])
        cell = table.row_cells(1)[10]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[2][0], to_replace[2][1])
        cell = table.row_cells(8)[0]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[3][0], to_replace[3][1])
        cell = table.row_cells(9)[0]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[4][0], to_replace[4][1])


    def clone_block(self, table: Table, index: int, count: int = 1):
        """Клонирует блок. 1 блок - 1 период

        Args:
            index (int): индекс первой строки копируемого блока
            count (int): количество копий.
        """

        stride = 4
        for i in range(count):
            for j in range(stride):
                new_row = self.redactor.insert_row_in_table(table, index + (i + 1) * stride + j)
                self.redactor.clone_table_row(table.rows[index + j], new_row)

            self.merge_row_5(table, index + (i + 1) * stride)
            self.merge_row_6(table, index + (i + 1) * stride + 1)
            self.merge_row_5(table, index + (i + 1) * stride + 2, True)
            self.merge_row_6(table, index + (i + 1) * stride + 3)


    def fill_period(self, period, period_index, table: Table, filled_rows: int):
        self.fill_common_period_info(period, period_index, table, filled_rows)

        """
        template_offset - это индекс, по которому находится строка-шаблон header-строки. Он вычисляется
        следующим образом. 4 - это первые четыре строки в таблице, они содержат информацию о договоре.
        filled_rows - количество строк, которые заполнены в таблице (в это число не входят первые четыре 
        строки в таблице, которые не относятся не к каким периодам)
        """
        template_offset_1 = 4 + filled_rows
        template_offset_2 = template_offset_1 + 1

        # Этот флаг отвечает за то какую строку мы заполняем
        first_row_flag = False

        current_filled_rows = 0
        for row in period["rows"]:
            paste_offset = template_offset_1 + current_filled_rows + 3    # +3 - это плюс три строки с шаблонами
            if first_row_flag:
                paste_offset -= 1
            # print(f"Обрабатываем строку: {row}")

            if self.type_of_row(row) == 1:
                if not first_row_flag:
                    self.fill_row_type_1(table, template_offset_1, row)
                    first_row_flag = True
                else:
                    self.clone_row_type_1(table, template_offset_1 + 2, paste_offset)
                    self.fill_row_type_1(table, paste_offset, row)
                    # template_offset_1 = paste_offset

            elif self.type_of_row(row) == 2:
                self.clone_row_type_2(table, template_offset_1 + 1, paste_offset)
                self.fill_row_type_2(table, paste_offset, row)
                # template_offset_2 = paste_offset

            current_filled_rows += 1
            # paste_offset += 1
            # self.redactor.save()

        # Здесь мы добавляем единицу, поскольку каждый блок заканчивается строкой "итого"
        current_filled_rows += 1
        # Удаляем строки с шаблонами
        self.redactor.delete_row_in_table(table, template_offset_1 + 1)
        self.redactor.delete_row_in_table(table, template_offset_1 + 1)
        # current_filled_rows += 2

        return current_filled_rows


    def fill_common_period_info(self, period, period_index: int, table: Table, filled_rows: int):
        # print(f"месяц: {period['period']}")
        to_replace = (
            (self.borders("период месяц.год"),              period["period"]),
            (self.borders("сумма пени за период"),          period["total"]),
        )

        cell = table.row_cells(4 + filled_rows)[0]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])

        cell = table.row_cells(4 + filled_rows + 3)[11]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[1][0], to_replace[1][1])


    def type_of_row(self, row: dict):
        """Возвращает тип строки. Если в строке есть текст, значит это строка-заголовок и возвращаем 1.
        Если текста нет, значит это строка в которой есть формула, ставка, пени и тд, возвращаем 2
        1 - строка-заголовок
        2 - строка-куча чисел

        Args:
            row (dict): строка row представлена в виде словаря с полями, в том числе есть поле "text".

        Returns:
            _type_: 1 или 2
        """
        if row["text"] is None:
            return 2
        else:
            return 1


    def count_row_by_types(self, rows):
        n_header_rows = 0
        n_number_rows = 0
        for row in rows:
            row_type = self.type_of_row(row)
            if row_type == 1:
                n_header_rows += 1
            elif row_type == 2:
                n_number_rows += 1
            else:
                print(f"Строка row={row} не соответствует ни одному формату, не ясно как её обрабатывать.")
                raise RuntimeError(f"Invalid row: {row}")
        return (n_header_rows, n_number_rows)


    def clone_row_type_1(self, table: Table, source_row_index: int, _target_row_index: int = None, merge_flag: bool = False):
        """Клонирует хедер-строку.

        Args:
            table (Table): Таблица, в которой происходят изменения
            source_row_index (int, optional): Индекс строки, которую копируем (она является шаблонной).
            target_row_index (int, optional): Индекс, по которому вставляем копию строки. По 
            умолчанию вставляется сразу после шаблонной строки.
            
            Update
            merge_flag - нужно ли присоединять строку к первому столбцу
        """
        target_row_index = source_row_index if _target_row_index is None else _target_row_index
        new_row = self.redactor.insert_row_in_table(table, target_row_index)
        self.redactor.clone_table_row(table.rows[source_row_index], new_row)
        self.merge_row_5(table, target_row_index, merge_flag)
        self.redactor.merge_table_cells(table.rows[target_row_index - 1].cells[0], table.rows[target_row_index].cells[0])


    def clone_row_type_2(self, table: Table, source_row_index: int, _target_row_index: int = None):
        """Клонирует строку с кучей численных данных.

        Args:
            table (Table): Таблица, в которой происходят изменения
            source_row_index (int, optional): Индекс строки, которую копируем (она является шаблонной).
            target_row_index (int, optional): Индекс, по которому вставляем копию строки. По 
            умолчанию вставляется сразу после шаблонной строки.
        """
        target_row_index = source_row_index if _target_row_index is None else _target_row_index
        new_row = self.redactor.insert_row_in_table(table, target_row_index)
        self.redactor.clone_table_row(table.rows[source_row_index], new_row)
        self.merge_row_6(table, target_row_index)
        self.redactor.merge_table_cells(table.rows[target_row_index - 1].cells[0], table.rows[target_row_index].cells[0])


    def fill_row_type_1(self, table: Table, row_index: int, row: dict):
        to_replace = (
            (self.borders("сумма начисления"),              row["debt"]),
            (self.borders("дата начала начисления"),        row["period"][0]),
            (self.borders("вставляемый текст"),             row["text"]),
        )

        cell = table.row_cells(row_index)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])
        cell = table.row_cells(row_index)[3]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[1][0], to_replace[1][1])
        cell = table.row_cells(row_index)[4]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[2][0], to_replace[2][1])


    def fill_row_type_2(self, table: Table, row_index: int, row: dict):
        # print(f"row={row}")
        to_replace = (
            (self.borders("сумма начисления"),              row["debt"]),
            (self.borders("дата начала начисления"),        row["period"][0]),
            (self.borders("дата конца начисления"),         row["period"][1]),
            (self.borders("кол-во дней начисления"),        str(row["period"][2])),
            (self.borders("ставка"),                        row["penalty_period_info"][0]),
            (self.borders("доля ставки"),                   row["penalty_period_info"][1]),
            (self.borders("формула"),                       row["formulae"]),
            (self.borders("пени"),                          row["penalty"]),
        )
        cell = table.row_cells(row_index)[1]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[0][0], to_replace[0][1])
        cell = table.row_cells(row_index)[3]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[1][0], to_replace[1][1])
        cell = table.row_cells(row_index)[4]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[2][0], to_replace[2][1])
        cell = table.row_cells(row_index)[5]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[3][0], to_replace[3][1])

        cell = table.row_cells(row_index)[6]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[4][0], to_replace[4][1])
        cell = table.row_cells(row_index)[8]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[5][0], to_replace[5][1])
        cell = table.row_cells(row_index)[9]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[6][0], to_replace[6][1])
        cell = table.row_cells(row_index)[11]
        self.redactor.replace_text_in_paragraph(cell.paragraphs[0], to_replace[7][0], to_replace[7][1])


    def borders(self, string: str, start: str = "/*", end: str = "*/"):
        return start + string + end


    def clone_table(self, count: int = 1):
        # Берём первую таблицу
        original_table = self.redactor.get_table(0)

        for i in range(count):
            # Получаем XML-элемент таблицы
            table_element = original_table._element

            # Глубокая копия XML-элемента таблицы (включая все внутренние элементы)
            cloned_table_element = deepcopy(table_element)

            # Находим родительский элемент (обычно это <w:body>)
            parent_element = table_element.getparent()

            # Определяем позицию исходной таблицы
            index_in_parent = parent_element.index(table_element)

            # Создаём пустой параграф
            new_paragraph = self.doc.add_paragraph()
            new_paragraph_element = new_paragraph._element

            # Создаём клонированную таблицу (обёртка для XML-элемента)
            # from docx.table import Table
            cloned_table = Table(cloned_table_element, self.doc)

            # Вставляем пустой параграф сразу после исходной таблицы
            parent_element.insert(index_in_parent + 1, new_paragraph_element)

            # Вставляем клонированную таблицу после параграфа
            parent_element.insert(index_in_parent + 2, cloned_table_element)


    def merge_row_1(self, table, row_index):
        # Объединяем ячейки первой строки
        row = table.rows[row_index]
        self.redactor.merge_table_cells(row.cells[0], row.cells[1])
        for i in range(3, 12):
            self.redactor.merge_table_cells(row.cells[2], row.cells[i])


    def merge_row_2(self, table, row_index):
        # Объединяем ячейки второй строки
        row = table.rows[row_index]
        self.redactor.merge_table_cells(row.cells[0], row.cells[1])
        for i in range(3, 7):
            self.redactor.merge_table_cells(row.cells[2], row.cells[i])
        for i in range(8, 10):
            self.redactor.merge_table_cells(row.cells[7], row.cells[i])
        self.redactor.merge_table_cells(row.cells[10], row.cells[11])


    def merge_row_3_4(self, table, row_3_index):
        # Объединяем ячейки третьей и четвертой строк
        row_3 = table.rows[row_3_index]
        row_4 = table.rows[row_3_index + 1]
        for row in [row_3, row_4]:
            self.redactor.merge_table_cells(row.cells[1], row.cells[2])
            self.redactor.merge_table_cells(row.cells[6], row.cells[7])
            self.redactor.merge_table_cells(row.cells[9], row.cells[10])

        for i in range(4, 6):
            self.redactor.merge_table_cells(row_3.cells[3], row_3.cells[i])

        for i in [0, 1, 6, 7, 8, 9, 10, 11]:
            self.redactor.merge_table_cells(row_3.cells[i], row_4.cells[i])


    def merge_row_5(self, table, row_index, first_cell_flag = False):
        """
        first_cell_flag - флаг сообщает, нужно ли присоединять первую колонку текущей строки
        к общей колонке "период". По умолчанию не нужно.
        """
        row = table.rows[row_index]
        self.redactor.merge_table_cells(row.cells[1], row.cells[2])
        for i in range(5, 12):
            self.redactor.merge_table_cells(row.cells[4], row.cells[i])

        if first_cell_flag:
            previous_row = table.rows[row_index - 1]
            self.redactor.merge_table_cells(previous_row.cells[0], row.cells[0])


    def merge_row_6(self, table, row_index):
        row = table.rows[row_index]
        self.redactor.merge_table_cells(row.cells[1], row.cells[2])
        self.redactor.merge_table_cells(row.cells[6], row.cells[7])
        self.redactor.merge_table_cells(row.cells[9], row.cells[10])

        previous_row = table.rows[row_index - 1]
        self.redactor.merge_table_cells(previous_row.cells[0], row.cells[0])


    def merge_row_8(self, table, row_index):
        row = table.rows[row_index]
        for i in range(1, 12):
            self.redactor.merge_table_cells(row.cells[0], row.cells[i])


    def clone_payment_row(self, table: Table, count: int = 1, source_row_index: int = 5, target_row_index: int = None):
        """Клонирует строку с данными платежа.

        Args:
            table (Table): Таблица, в которой происходят изменения
            source_row_index (int, optional): Индекс строки, которую копируем (она является шаблонной). 
            По умолчанию = 5.
            target_row_index (int, optional): Индекс, по которому вставляем копию строки. По 
            умолчанию вставляется сразу после шаблонной строки.
        """

        _target_row_index = source_row_index + 1 if target_row_index is None else target_row_index

        for i in range(count):
            new_row = self.redactor.insert_row_in_table(table, _target_row_index)
            self.redactor.clone_table_row(table.rows[source_row_index], new_row)
            self.merge_row_6(table, _target_row_index)
            self.redactor.merge_table_cells(table.rows[_target_row_index - 1].cells[0], table.rows[_target_row_index].cells[0])


    def fill_second_table(self):
        self.fill_second_table_common_info()

        table = self.redactor.get_table(1)
        # self.redactor.print_table(table)
        
        # по сути число договоров
        n_contracts = len(self.config2["contracts_info"])

        # По смыслу, это количство уже заполненных строк в таблице. Начинаем не с 0, а с 1 потому что
        # в таблице есть заголовочная строка, в которой ничего заполнять не нужна, но она есть сама 
        # по себе и её нужно учитывать. По сути, значение n_row равно индексу строки, которую мы будем 
        # заполнять на текущей итерации
        n_rows = 1
        for i, contract in enumerate(self.config2["contracts_info"]):

            if i < n_contracts - 1:
                self.second_table_clone_row(table, n_rows)

            contract_number = list(self.config2["table_info"].keys())[i]

            if self.config2["table_info"][contract_number]["correcting_debt"] == "0,00":
                self.fill_second_table_simple_row(table, n_rows, i)
                n_rows += 1

            else:
                self.second_table_clone_row(table, n_rows)
                self.fill_second_table_complex_row(table, n_rows, i)
                n_rows += 2


    def fill_second_table_common_info(self):
        table = self.redactor.get_table(1)
        self.redactor.replace_text_in_paragraph(
            table.row_cells(2)[2].paragraphs[0],
            self.borders("сумма долга"),
            self.config2["table_info"]["all_debt"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(2)[3].paragraphs[0],
            self.borders("неустойка общая"),
            self.config2["table_info"]["all_penalty"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(2)[4].paragraphs[0],
            self.borders("цена иска"),
            self.config2["table_info"]["cost_of_lawsuit"]
        )


    def second_table_clone_row(self, table: Table, cloning_row_index: int):
        new_row = self.redactor.insert_row_in_table(table, cloning_row_index + 1)
        self.redactor.clone_table_row(table.rows[cloning_row_index], new_row)


    def fill_second_table_simple_row(self, table: Table, row_index: int, contract_index: int):
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[0].paragraphs[0],
            self.borders("номер договора"),
            self.config2["contracts_info"][contract_index][1]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[1].paragraphs[0],
            self.borders("период"),
            self.config2["contracts_info"][contract_index][2]["contract_periods"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config2["contracts_info"][contract_index][2]["debt"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[3].paragraphs[0],
            self.borders("неустойка"),
            # "#неустойка#".upper()
            self.config2["contracts_info"][contract_index][2]["penalty"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[4].paragraphs[0],
            self.borders("неустойка+задолженность"),
            # "#неустойка+задолженность#".upper()
            self.config2["contracts_info"][contract_index][2]["debt_penalty"]
        )


    def fill_second_table_complex_row(self, table: Table, row_index: int, contract_index: int):
        self.second_table_merge_rows(table, row_index)

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[0].paragraphs[0],
            self.borders("номер договора"),
            self.config2["contracts_info"][contract_index][1]
        )

        # Ячейка Период
        source_paragraph = table.row_cells(row_index)[1].paragraphs[0]
        new_paragraph = table.row_cells(row_index)[1].add_paragraph("")
        self.redactor.clone_paragraph(source_paragraph, new_paragraph)
        self.redactor.replace_text_in_paragraph(
            source_paragraph,
            self.borders("период"),
            self.config2["contracts_info"][contract_index][2]["contract_periods"]
        )
        self.redactor.replace_text_in_paragraph(
            new_paragraph,
            self.borders("период"),
            "текущие начисления"
        )
        self.redactor.paragraph_text_set_bold(new_paragraph)

        source_paragraph = table.row_cells(row_index + 1)[1].paragraphs[0]
        new_paragraph = table.row_cells(row_index + 1)[1].add_paragraph("")
        self.redactor.clone_paragraph(source_paragraph, new_paragraph)
        self.redactor.replace_text_in_paragraph(
            source_paragraph,
            self.borders("период"),
            self.config2["contracts_info"][contract_index][2]["contract_periods_correcting"]
        )
        self.redactor.replace_text_in_paragraph(
            new_paragraph,
            self.borders("период"),
            "доля от ГК"
        )
        self.redactor.paragraph_text_set_bold(new_paragraph)

        # Ячейка Задолженность
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config2["contracts_info"][contract_index][2]["debt"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index + 1)[2].paragraphs[0],
            self.borders("задолженность"),
            self.config2["contracts_info"][contract_index][2]["correcting_debt"]
        )

        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[3].paragraphs[0],
            self.borders("неустойка"),
            # "#неустойка#".upper()
            self.config2["contracts_info"][contract_index][2]["penalty"]
        )
        self.redactor.replace_text_in_paragraph(
            table.row_cells(row_index)[4].paragraphs[0],
            self.borders("неустойка+задолженность"),
            # "#неустойка+задолженность#".upper()
            self.config2["contracts_info"][contract_index][2]["debt_penalty"]
        )


    def second_table_merge_rows(self, table: Table, row_index: int):
        row_1 = table.rows[row_index]
        row_2 = table.rows[row_index + 1]
        for i in [0, 3, 4]:
            self.redactor.merge_table_cells(row_1.cells[i], row_2.cells[i])


    def fill_other_parts(self):
        to_replace = {
            "/*цена иска*/": self.config2["lawsuit_info"]["cost"],
            "/*госпошлина*/": self.config2["lawsuit_info"]["tax"],
            "/*текущая дата*/": datetime.strptime(self.config2["current_date"], "%Y-%m-%d").strftime("%d.%m.%Y")
        }

        for paragraph in self.doc.paragraphs:
            # runs = []
            for run in paragraph.runs:
                # runs.append(run.text)
                for k, v in to_replace.items():
                    if k in run.text:
                        run.text = run.text.replace(k, v)
            # print(f"runs={runs}")
