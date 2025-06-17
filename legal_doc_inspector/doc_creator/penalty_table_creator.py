from collections import defaultdict
from pathlib import Path
from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.table import Table, _Cell, _Row, _Column
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# from .utils.month_to_decimal import month_year_to_decimal
from legal_doc_inspector.doc_creator.utils.month_to_decimal import month_year_to_decimal

class PenaltyTableCreator:
    
    def __init__(self):
        self.doc = Document()
        self._create_document_title()
        self.contracts_info = defaultdict(lambda : {})


    def save_doc(self, path_to_save:Path):
        self.doc.save(str(path_to_save))
    
    def create_penalty_table_from_json(self, name, data, contract_number, start_date, end_date):
        
        table = self._create_table_title(contract_number, start_date, end_date)
        list_of_periods = self.group_by_month(data)
        table , contract_number, all_debt, all_penalty = self._create_penalty_table(table, list_of_periods,contract_number)

        return contract_number, start_date, end_date, all_debt, all_penalty


    def create_result_table(self, list_of_tables_info:list[tuple], name:Path):

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('\n \n \nII.   РАСЧЕТ ОБЩЕЙ ЦЕНЫ ИСКА')
        run.font.name = 'Times New Roman'
        run.bold = True
        element = run._element
        rPr = element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')  
        rFonts.set(qn('w:cs'), 'Times New Roman')
        run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        result_json_part = defaultdict(lambda : {})
        table = self.doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        title_row_cells = table.rows[0].cells
        self._put_text_into_table_cell(text='Реквизиты (номер) договора',
                                       cell=title_row_cells[0],
                                       font_size=12,
                                       need_bold=True)
        
        self._put_text_into_table_cell(text='Период',
                                       cell=title_row_cells[1],
                                       font_size=12,
                                       need_bold=True)
        
        self._put_text_into_table_cell(text='Задолженность (основной долг), руб.',
                                       cell=title_row_cells[2],
                                       font_size=12,
                                       need_bold=True)
        
        self._put_text_into_table_cell(text='Неустойка,руб.',
                                       cell=title_row_cells[3],
                                       font_size=12,
                                       need_bold=True)
        
        self._put_text_into_table_cell(text='Итого по договору(неустойка + основной долг), сумма, руб.',
                                       cell=title_row_cells[4],
                                       font_size=12,
                                       need_bold=True)
        all_debt = 0
        all_penalty = 0

        for contract_number, start_date, end_date, debt, penalty in list_of_tables_info:
            
            self.contracts_info[contract_number]['penalty_period'] = f'{start_date} по {end_date}'
            row = table.add_row()
            row_cells = row.cells

            self._put_text_into_table_cell(text=f'{contract_number}',
                                           font_size=11,
                                           cell=row_cells[0],
                                           need_bold=True)
            
            self._put_text_into_table_cell(text=f"{self.contracts_info[contract_number]['contract_periods']}",
                                           font_size=11,
                                           cell=row_cells[1])
            
            self._put_text_into_table_cell(text=f'{self.format_float_to_currency(debt)}',
                                           cell=row_cells[2],
                                           font_size=11)
            
            self._put_text_into_table_cell(text=f'{self.format_float_to_currency(penalty)}',
                                           cell=row_cells[3],
                                           font_size=11)
            
            self._put_text_into_table_cell(text=f'{self.format_float_to_currency(float(debt) + float(penalty))}',
                                           cell=row_cells[4],
                                           font_size=11)
            
            self.contracts_info[contract_number]['debt'] = self.format_float_to_currency(debt)
            self.contracts_info[contract_number]['penalty'] = self.format_float_to_currency(penalty)
            self.contracts_info[contract_number]['debt_penalty'] = self.format_float_to_currency(float(debt) + float(penalty))

            all_debt+=float(debt)
            all_penalty+=float(penalty)

        result_row = table.add_row()

        result_row_cells = result_row.cells

        result_row_cells[0].merge(result_row_cells[1])

        self._put_text_into_table_cell(text="ИТОГО",
                                       font_size=11,
                                       cell=result_row_cells[0],
                                       need_bold=True,
                                       orient='right')
        
        self._put_text_into_table_cell(text=f'{self.format_float_to_currency(all_debt)}',
                                       font_size=11,
                                       need_bold=True,
                                       cell=result_row_cells[2])
        
        self._put_text_into_table_cell(text=f'{self.format_float_to_currency(all_penalty)}',
                                       font_size=11,
                                       need_bold=True,
                                       cell=result_row_cells[3])
        
        self._put_text_into_table_cell(text=f'{self.format_float_to_currency(float(all_debt) + float(all_penalty))}',
                                       font_size=11,
                                       need_bold=True,
                                       cell=result_row_cells[4])

        self.save_doc(name)
        self.contracts_info['all_debt'] = self.format_float_to_currency(all_debt)
        self.contracts_info['all_penalty'] = self.format_float_to_currency(all_penalty)
        self.contracts_info['cost_of_lawsuit'] = self.format_float_to_currency(float(all_debt) + float(all_penalty))

        return str(name.resolve()), self.contracts_info

    def _create_document_title(self):

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('Приложение к исковому заявлению')
        run.font.name = 'Times New Roman'
        element = run._element
        rPr = element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')  
        rFonts.set(qn('w:cs'), 'Times New Roman')
        run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('РАСЧЕТ СУММЫ НЕУСТОЙКИ (ПЕНЕЙ) И ОБЩЕЙ ЦЕНЫ ИСКА')
        run.font.name = 'Times New Roman'
        run.bold = True
        element = run._element
        rPr = element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')  
        rFonts.set(qn('w:cs'), 'Times New Roman')
        run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run('I.   РАСЧЕТ СУММЫ НЕУСТОЙКИ (ПЕНЕЙ)')
        run.font.name = 'Times New Roman'
        run.bold = True
        element = run._element
        rPr = element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')  
        rFonts.set(qn('w:cs'), 'Times New Roman')
        run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    def _create_table_title(self, contract_number:str, start_date, end_date ):
        
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(f'Исходя из размера Основного долга и Периода задолженности, неустойка по состоянию на {end_date} составляет:')
        run.font.name = 'Times New Roman'
        element = run._element
        rPr = element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')  
        rFonts.set(qn('w:cs'), 'Times New Roman')
        run.font.size = Pt(12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc = self.doc
        
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Pt(842)  # ширина
        section.page_height = Pt(595)  # высота

        # создание титульника таблицы

        table = doc.add_table(rows=1, cols=11)
        table.style = 'Table Grid'
        self._apply_height_for_row(table.rows[0],0.6)
        row_cells = table.rows[0].cells
        
        row_cells[0].merge(row_cells[2])

        self._put_text_into_table_cell('Информация о расчёте',
                                        row_cells[0],
                                        need_gray_bgc=True,
                                        orient='left')

        row_cells[3].merge(row_cells[10])
    

        self._put_text_into_table_cell(text=f'{contract_number}',
                                       cell=row_cells[3],
                                       orient='left',
                                       need_bold=True)

        secondrow = table.add_row()
        self._apply_height_for_row(secondrow, 0.6)
        row_cells = secondrow.cells
        row_cells[0].merge(row_cells[1])

        self._put_text_into_table_cell(text="Начало просрочки",
                                       cell=row_cells[0],
                                       need_gray_bgc=True,
                                       orient='left'
                                       )
        row_cells[2].merge(row_cells[4])

        self._put_text_into_table_cell(text=f'{start_date}',
                                       cell=row_cells[2],
                                       orient='left')

        row_cells[5].merge(row_cells[6])

        self._put_text_into_table_cell(text="конец просрочки",
                                       cell=row_cells[5],
                                       orient='left',
                                       need_gray_bgc=True)

        row_cells[7].merge(row_cells[10])

        self._put_text_into_table_cell(text=f'{end_date}',
                                       cell=row_cells[7],
                                       orient='left')

        third_row = table.add_row()
        self._apply_height_for_row(third_row, 0.6)
        fourth_row = table.add_row()
        self._apply_height_for_row(fourth_row, 0.6)
        up_row_cells = third_row.cells
        down_row_cells = fourth_row.cells

        up_row_cells[0].merge(up_row_cells[1])
        down_row_cells[0].merge(down_row_cells[1])
        up_row_cells[0].merge(down_row_cells[0])

        self._put_text_into_table_cell(text='Месяц',
                                       need_bold=True,
                                       cell=up_row_cells[0])

        up_row_cells[2].merge(up_row_cells[3])
        down_row_cells[2].merge(down_row_cells[3])
        up_row_cells[2].merge(down_row_cells[2])

        self._put_text_into_table_cell(text='Долг',
                                       need_bold=True,
                                       cell=up_row_cells[2])

        up_row_cells[4].merge(up_row_cells[6])

        self._put_text_into_table_cell(text='Период просрочки',
                                       cell=up_row_cells[4],
                                       need_bold=True)


        self._put_text_into_table_cell(text='с',
                                       need_bold=True,
                                       cell=down_row_cells[4])

        self._put_text_into_table_cell(text='по',
                                       need_bold=True,
                                       cell=down_row_cells[5])

        self._put_text_into_table_cell(text='дней',
                                       need_bold=True,
                                       cell=down_row_cells[6])


        up_row_cells[7].merge(down_row_cells[7])

        self._put_text_into_table_cell(text='Ставка',
                                       need_bold=True,
                                       cell=up_row_cells[7])

        up_row_cells[8].merge(down_row_cells[8])

        self._put_text_into_table_cell(text='Доля ставки',
                                       need_bold=True,
                                       cell=up_row_cells[8])

        up_row_cells[9].merge(down_row_cells[9])

        self._put_text_into_table_cell(text="Формула",
                                       need_bold=True,
                                       cell=up_row_cells[9])

        up_row_cells[10].merge(down_row_cells[10])

        self._put_text_into_table_cell(text="Пени",
                                       need_bold=True,
                                       cell=up_row_cells[10])
        
        return table
    

    def group_by_month(self, data):
        """
        принимает результат калькулятора пени, 
        возвращает периоды оплат и просрочек для каждого месяца
        """

        grouped = defaultdict(lambda: {'payments': [], 'periods': []})

        # Группируем платежи и периоды по месяцу
        for entry in data:
            month_key = entry['month']
            if 'payments' in entry:
                grouped[month_key]['payments'].append(entry['payments'])
            # Убираем лишние поля перед добавлением периода
            period_data = {
                k: v for k, v in entry.items()
                if k not in ('payments', 'month')
            }
            grouped[month_key]['periods'].append(period_data)

        result = []
        for month, items in grouped.items():
            combined = []
            # Сначала добавляем все платежи
            payments_seen = set()
            
            for payment in items['payments']:
                if payment is None:
                    continue
                
                for key, value in payment.items():
                    if key not in payments_seen:
                        # print(key)
                        payments_seen.add(key)
                        combined.append({'type': 'payment', 'data': {key:value}})
                
                
            # Затем добавляем все периоды
            for period in items['periods']:
                combined.append({'type': 'period', 'data': period})

            def get_sort_key(event):
                if event['type'] == 'payment':
                    # Берём единственную дату из ключей
                    return next(iter(event['data'].keys()))
                else:
                    return event['data']['start']
                
            combined.sort(key=get_sort_key)

            result.append({month: combined})


        return result

    
    def _create_penalty_table(self, table:Table, list_dict_of_months, contract_number):
        """
        принимает таблицу в документе и  результат group by month,
        добавляет в таблицу строки периодов для каждого месяца
        """
        all_debt = 0
        all_penalty = 0
        months = []
        for month_dict in list_dict_of_months:
            month_key, periods_and_payments  = next(iter(month_dict.items()))
            months.append(month_key)
            new_rows_for_month = []
            itogo = 0
            last_month_debt = 0
            for period_or_payment in periods_and_payments:
                if period_or_payment['type'] == 'period':
                    data = period_or_payment['data']
                    new_row = table.add_row()
                    self._apply_height_for_row(new_row, 0.65)
                    new_row_cells = new_row.cells
                    debt = data['debt']
                    last_month_debt = debt
                    start_date = data['start']
                    end_date = data['end']
                    days_count = data['days']
                    rate = data['rate']
                    rate = self._rate_float_to_decimal(rate)
                    penalty = data['penalty']
                    itogo+=penalty

                    start_date = start_date.strftime("%d.%m.%Y")
                    end_date = end_date.strftime("%d.%m.%Y")

                    new_row_cells[0].merge(new_row_cells[1])
                    new_row_cells[2].merge(new_row_cells[3])

                    
                    self._put_text_into_table_cell(text=f'{self.format_float_to_currency(debt)}',
                                                   cell=new_row_cells[2],)

                    self._put_text_into_table_cell(text=f'{start_date}',
                                                   cell=new_row_cells[4] )

                    self._put_text_into_table_cell(text=f'{end_date}',
                                                   cell=new_row_cells[5])
                    
                    self._put_text_into_table_cell(text=f'{days_count}',
                                                   cell=new_row_cells[6])
                    
                    self._put_text_into_table_cell(text='9,50%',
                                                   cell=new_row_cells[7])

                    self._put_text_into_table_cell(text=rate,
                                                   cell=new_row_cells[8])
                    
                    self._put_text_into_table_cell(text=f'{self.format_float_to_currency(debt)} × {days_count} × {rate} × 9,5%',
                                                   cell=new_row_cells[9])

                    self._put_text_into_table_cell(text=f'{self.format_float_to_currency(penalty)} р.',
                                                   cell=new_row_cells[10],
                                                   orient='right')

                    new_rows_for_month.append(new_row)

                if period_or_payment['type'] == 'payment':
                    data = period_or_payment['data']
                    new_row = table.add_row()
                    self._apply_height_for_row(new_row, 0.65)

                    new_row_cells = new_row.cells
                    date_of_payment, amount = next(iter(data.items()))
                    date_of_payment = date_of_payment.strftime("%d.%m.%Y")

                    new_row_cells[0].merge(new_row_cells[1])
                    new_row_cells[2].merge(new_row_cells[3])
                    self._put_text_into_table_cell(text=f'-{self.format_float_to_currency(amount)}',
                                                   cell=new_row_cells[2])

                    self._put_text_into_table_cell(text=f'{date_of_payment}',
                                                   cell=new_row_cells[4])
                    
                    self._put_text_into_table_cell(text="Погашение части долга",
                                                   need_italic=True,
                                                   cell=new_row_cells[5],
                                                   orient='left')
                    
                    new_row_cells[5].merge(new_row_cells[10])
                    new_rows_for_month.append(new_row)
            
            itog_row = table.add_row()
            itog_row_cells = itog_row.cells
            itog_row_cells[0].merge(itog_row_cells[1])
            itog_row_cells[2].merge(itog_row_cells[3])
            self._put_text_into_table_cell(text="Итого",
                                           cell=itog_row_cells[9],
                                           need_bold=True,
                                           need_gray_bgc=True,
                                           orient='right')
            
            self._put_text_into_table_cell(text=f'{self.format_float_to_currency(itogo)} р.',
                                           cell=itog_row_cells[10],
                                           need_bold=True,
                                           need_gray_bgc=True,
                                           orient='right')
            
            new_rows_for_month.append(itog_row)
            new_rows_for_month[0].cells[0].merge(new_rows_for_month[-1].cells[0])
            
            self._put_text_into_table_cell(text=f'{month_key}',
                                           cell=new_rows_for_month[0].cells[0],
                                           need_bold=True,
                                           need_vertical_orient=True)


            self._set_row_bottom_border(new_rows_for_month[-1])

            self._apply_width_for_column(table.columns[9], 5)
            self._apply_width_for_column(table.columns[2], 0.7)
            

            all_penalty+=itogo
            all_debt += last_month_debt

        self.contracts_info[contract_number]['contract_periods'] = (f'{month_year_to_decimal(months[0])}-{month_year_to_decimal(months[-1])}')
        first_row = table.add_row()
        second_row = table.add_row()
        self._apply_height_for_row(first_row, 0.65)
        self._apply_height_for_row(second_row, 0.65)
        
        first_row_cells = first_row.cells
        first_row_cells[0].merge(first_row_cells[-1])
        self._put_text_into_table_cell(text=f'Сумма основного долга: {self.format_float_to_currency(all_debt)} руб.',
                                       cell=first_row_cells[0],
                                       orient='right')
        
        second_row_cells = second_row.cells
        second_row_cells[0].merge(second_row_cells[-1])
        self._put_text_into_table_cell(text=f'Сумма пеней по всем задолженностям: {self.format_float_to_currency(all_penalty)} руб.',
                                       cell=second_row_cells[0],
                                       orient='right')

        return table, contract_number, all_debt, all_penalty

    

    def _rate_float_to_decimal(self,rate):
        """
        переводит число с плавающей точкой в обыкновенную дробь
        """
        first = 1/300
        second = 1/170
        third = 1/130
        fourth = 0
        
        if rate==first:    
            return "1/300"
        if rate==second:
            return "1/170"
        if rate==third:
            return "1/130"
        if rate==fourth:
            return "0"
        

    

    def convert_datetime_keys(self, obj):
        if isinstance(obj, dict):
            new_dict = {}
            for k, v in obj.items():
                # Преобразуем ключ, если он datetime/date
                new_key = k
                if isinstance(k, date):
                    new_key = k.strftime("%d.%m.%Y")
                
                new_value = self.convert_datetime_keys(v)
                new_dict[new_key] = new_value
            return new_dict
        elif isinstance(obj, list):
            return [self.convert_datetime_keys(item) for item in obj]
        elif isinstance(obj, date):
            return obj.strftime("%d.%m.%Y")
        else:
            return obj
        

    def _put_text_into_table_cell(self, text:str, cell:_Cell, font_size=9, need_bold=False, need_italic=False, orient="center", need_vertical_orient=True, need_gray_bgc=False):
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(text)
        run.bold = need_bold
        run.italic = need_italic
        run.font.name = 'Times New Roman'

        element = run._element
        rPr = element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')  
        rFonts.set(qn('w:cs'), 'Times New Roman')
        # run._element.rPr.rFonts.set('w:eastAsia', 'Times New Roman')  # для корректного отображения в Word
        run.font.size = Pt(font_size)

        match orient:
            case "center":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            case "left":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            case "right":
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if need_vertical_orient:
            self._set_cell_vertical_alignment(cell=cell)

        if need_gray_bgc:
            self._set_cell_background(cell=cell)


    def _set_cell_vertical_alignment(self, cell:_Cell, align="center"):
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tag = tc_pr.xpath('w:vAlign')
        if tag:
            el = tag[0]
            el.set(qn('w:val'), align)
        else:
            valign = OxmlElement('w:vAlign')
            valign.set(qn('w:val'), align)
            tc_pr.append(valign)


    def _set_cell_background(self, cell:_Cell, color_hex = "#cccccc"):
        tc_pr = cell._tc.get_or_add_tcPr()
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        tc_pr.append(shading_elm)


    def _set_row_bottom_border(self, row:_Row, color="000000", size=4):
        """ 
        :param row: объект строки (Row)
        :param color: цвет в HEX без решётки (#), например "000000" — чёрный
        :param size: толщина линии (в 1/8 pt, например 4 = 0.5pt, 6 = 0.75pt)
        """
        tr = row._tr
        tc_borders = OxmlElement('w:tblBottomBorders')

        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')
        bottom_border.set(qn('w:sz'), str(size))
        bottom_border.set(qn('w:color'), color)  

        tc_borders.append(bottom_border)

        for el in tr.xpath('w:tblBottomBorders'):
            tr.remove(el)

        tr.append(tc_borders)
    

    def _apply_height_for_row(self, row:_Row, height:float):
        row.height = Cm(height)
        
    def format_float_to_currency(self, value: float) -> str:
        value = round(value, 2)  # избегаем ошибок округления
        rubbles = int(value)
        kopecks = int(round(value - rubbles, 2) * 100)

        # Форматируем рубли
        rubbles_str = str(rubbles)
        formatted = ''
        for i, digit in enumerate(rubbles_str[::-1]):
            if i % 3 == 0 and i != 0:
                formatted = ' ' + formatted
            formatted = digit + formatted

        return f"{formatted},{kopecks:02d}"
        
    
    def _apply_width_for_column(self, column:_Column, width:float):
        column.width = Cm(width)