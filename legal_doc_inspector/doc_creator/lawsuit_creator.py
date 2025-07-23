from pathlib import Path
import datetime

import docx
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.table import Table, _Cell, _Row, _Column
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns

from docx import Document


nbsp_replacements = {
    "каб ": "каб\u00a0",
    "ул. ": "ул.\u00a0",
    "д. ": "д.\u00a0",
    "к. ": "к.\u00a0",
    "стр. ": "стр.\u00a0",
    "помещ. ": "помещ.\u00a0",
    "г. ": "г.\u00a0",
    "ОГРН ": "ОГРН\u00a0",
    "ИНН ": "ИНН\u00a0",
    "офис " :"офис\u00a0",
    "аал ": "аал\u00a0",
    "аллея ": "аллея\u00a0",
    "аул ": "аул\u00a0",
    "б-р ": "б-р\u00a0",
    "версаль ": "версаль\u00a0",
    "высел ": "высел\u00a0",
    "городок ": "городок\u00a0",
    "д ": "д\u00a0",
    "дор ": "дор\u00a0",
    "ж/д_будка ": "ж/д_будка\u00a0",
    "ж/д_казарм ": "ж/д_казарм\u00a0",
    "ж/д_ст ": "ж/д_ст\u00a0",
    "ж/д_пост ": "ж/д_пост\u00a0",
    "ж/д_рзд ": "ж/д_рзд\u00a0",
    "ж/д_оп ": "ж/д_оп\u00a0",
    "заезд ": "заезд\u00a0",
    "казарма ": "казарма\u00a0",
    "кв-л ": "кв-л\u00a0",
    "км ": "км\u00a0",
    "кольцо ": "кольцо\u00a0",
    "линия ": "линия\u00a0",
    "м ": "м\u00a0",
    "мкр ": "мкр\u00a0",
    "наб ": "наб\u00a0",
    "нп ": "нп\u00a0",
    "остров ": "остров\u00a0",
    "парк ": "парк\u00a0",
    "переезд ": "переезд\u00a0",
    "пер ": "пер\u00a0",
    "пл-ка ": "пл-ка\u00a0",
    "пл ": "пл\u00a0",
    "платф ": "платф\u00a0",
    "полустанок ": "полустанок\u00a0",
    "п/ст ": "п/ст\u00a0",
    "п ": "п\u00a0",
    "починок ": "починок\u00a0",
    "п/о ": "п/о\u00a0",
    "п/р ": "п/р\u00a0",
    "просек ": "просек\u00a0",
    "проселок ": "проселок\u00a0",
    "пр-кт ": "пр-кт\u00a0",
    "проезд ": "проезд\u00a0",
    "проулок ": "проулок\u00a0",
    "рзд ": "рзд\u00a0",
    "с ": "с\u00a0",
    "сад ": "сад\u00a0",
    "сквер ": "сквер\u00a0",
    "сл ": "сл\u00a0",
    "ст ": "ст\u00a0",
    "стр ": "стр\u00a0",
    "тер ": "тер\u00a0",
    "тракт ": "тракт\u00a0",
    "туп ": "туп\u00a0",
    "ул ": "ул\u00a0",
    "уч-к ": "уч-к\u00a0",
    "х ": "х\u00a0",
    "ш ": "ш\u00a0"
}

class LawsuitCreator:

    def __init__(self, data_json:dict):
        pass

    def create_lawsuit(self, lawsuit_json, result_file_path):
        replace_list = [
            ("/*цена иска*/", lawsuit_json["lawsuit_info"]["cost"]),
            ("/*госпошлина*/", lawsuit_json["lawsuit_info"]["tax"]),
            ("/*ответчик*/", lawsuit_json["defendant_info"]["short_name"]),
            ("/*инн*/", lawsuit_json["defendant_info"]["inn"]),
            ("/*огрн*/", lawsuit_json["defendant_info"]["ogrn"]),
            ("/*номер договора*/", lawsuit_json["contracts_info"][0][1]),
            ("/*период*/", lawsuit_json["contracts_info"][0][2]["contract_periods"]),
            ("/*задолженность*/", lawsuit_json[]),
            ("/*неустойка*/", lawsuit_json[]),
            ("/*неустойка+задолженность*/", lawsuit_json[]),
            ("/*срок оплаты*/", lawsuit_json["contracts_info"][0][2]["last_day"]),
            ("/*пункт*/", lawsuit_json["contracts_info"][0][2]["contract_point"]),
            ("/*сумма долга*/", lawsuit_json[]),
            ("/*тип договора*/", lawsuit_json["lawsuit_info"]["service_type"]),
            ("/*номер претензии*/", lawsuit_json["lawsuit_info"]["claims"][0]),
            ("/*неустойка общая*/", lawsuit_json[]),
        ]
        self.__replace_text_preserve_formatting("template.docx", replace_list, result_file_path)


    def __replace_text_preserve_formatting(docx_path, replaces, output_path):
        """
        Replace text in a DOCX document while preserving formatting.

        Args:
            docx_path (str): Path to the input DOCX file
            old_text (str): Text to be replaced
            new_text (str): New text to insert
            output_path (str): Path to save the modified document
        """
        doc = Document(docx_path)

        for r in replaces:
            old_text, new_text = r
            for paragraph in doc.paragraphs:
                if old_text in paragraph.text:
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        if old_text in inline[i].text:
                            text = inline[i].text.replace(old_text, new_text)
                            inline[i].text = text

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if old_text in paragraph.text:
                                inline = paragraph.runs
                                for i in range(len(inline)):
                                    if old_text in inline[i].text:
                                        text = inline[i].text.replace(old_text, new_text)
                                        inline[i].text = text

        doc.save(output_path)
